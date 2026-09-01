from pathlib import Path

import pandas as pd
from docx import Document

from .profiling import profile_file, read_tables


SPREADSHEET_EXTENSIONS = {".xlsx", ".xls", ".csv", ".tsv", ".json"}
DOCUMENT_EXTENSIONS = {".docx"}


def supported(filename: str) -> bool:
    return Path(filename).suffix.lower() in SPREADSHEET_EXTENSIONS | DOCUMENT_EXTENSIONS


def inspect_file(path: Path) -> dict:
    suffix = path.suffix.lower()
    if suffix in SPREADSHEET_EXTENSIONS:
        profile = profile_file(path)
        first_table = max(
            profile.tables,
            key=lambda table: (table.row_count * len(table.columns), table.row_count, len(table.columns)),
            default=None,
        )
        return {
            "kind": "spreadsheet",
            "rows": first_table.row_count if first_table else 0,
            "columns": [column.name for column in first_table.columns] if first_table else [],
            "missing_cells": first_table.missing_cells if first_table else 0,
            "profile": profile.to_dict(),
        }
    document = Document(path)
    statements = [paragraph.text.strip() for paragraph in document.paragraphs if paragraph.text.strip()]
    return {
        "kind": "document",
        "paragraph_count": len(statements),
        "text_preview": statements[:5],
    }


def read_spreadsheet(path: Path) -> pd.DataFrame:
    if path.suffix.lower() == ".csv":
        return pd.read_csv(path)
    if path.suffix.lower() == ".tsv":
        return pd.read_csv(path, sep="\t")
    if path.suffix.lower() == ".json":
        tables = read_tables(path)
        if not tables:
            return pd.DataFrame()
        frame = tables[0][1].copy()
        frame.attrs["source_sheet"] = tables[0][0]
        frame.attrs["header_row"] = 0
        return frame
    return _best_table(pd.ExcelFile(path))


def _best_table(workbook: pd.ExcelFile) -> pd.DataFrame:
    candidates = []
    for sheet_name in workbook.sheet_names:
        raw = pd.read_excel(workbook, sheet_name=sheet_name, header=None)
        for header_row in range(min(len(raw), 12)):
            header = raw.iloc[header_row]
            columns = [index for index, value in header.items() if isinstance(value, str) and value.strip()]
            labels = [str(header[index]).strip() for index in columns]
            if len(labels) < 2 or len(labels) != len(set(labels)):
                continue
            data = raw.iloc[header_row + 1 :, columns].dropna(how="all")
            if data.empty:
                continue
            frame = data.copy()
            frame.columns = labels
            frame = frame.infer_objects(copy=False)
            time_columns = [
                label
                for label in labels
                if not pd.api.types.is_numeric_dtype(frame[label])
                and pd.to_datetime(frame[label], format="mixed", errors="coerce").notna().sum() >= 2
            ]
            metric_columns = [label for label in labels if label not in time_columns and pd.to_numeric(frame[label], errors="coerce").notna().sum() >= 2]
            valid_rows = pd.Series(False, index=frame.index)
            for time_column in time_columns:
                for metric_column in metric_columns:
                    valid_rows |= pd.to_datetime(frame[time_column], format="mixed", errors="coerce").notna() & pd.to_numeric(frame[metric_column], errors="coerce").notna()
            if not time_columns or not metric_columns or valid_rows.sum() < 2:
                continue
            score = len(labels) * 100 + len(time_columns) * 20 + len(metric_columns) * 20 + int(valid_rows.sum())
            candidates.append((score, sheet_name, header_row, labels, frame))

    if not candidates:
        frame = pd.read_excel(workbook, sheet_name=workbook.sheet_names[0])
        frame.attrs["source_sheet"] = workbook.sheet_names[0]
        frame.attrs["header_row"] = 0
        return frame

    _, sheet_name, header_row, labels, data = max(candidates, key=lambda candidate: candidate[0])
    frame = data.copy().infer_objects(copy=False)
    frame.columns = labels
    frame = frame.reset_index(drop=True)
    frame.attrs["source_sheet"] = sheet_name
    frame.attrs["header_row"] = header_row
    return frame

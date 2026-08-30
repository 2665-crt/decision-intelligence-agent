from html import escape
from pathlib import Path

from docx import Document


def render(job: dict, result: dict, directory: Path, route: str = "jobs") -> list[dict]:
    reports = directory / "reports"
    reports.mkdir(exist_ok=True)
    sections = result.get("sections", [])
    risks = result.get("business_risks", result.get("risks", []))
    suggestions = result.get("suggestions", result.get("options", []))
    quality = result.get("data_quality", {})
    lines = [
        "# 数据分析决策报告",
        f"\n## 分析目标\n{job['objective']}",
        f"\n## 核心结论\n{result.get('core_conclusion', '未生成核心结论。')}",
        "\n## 关键数据",
        *[f"- {item['label']}：{item['value']}。{item.get('detail', '')}" for item in result.get("key_metrics", [])],
        "\n## 详细分析",
    ]
    for section in sections:
        lines.extend([f"\n### {section['title']}", *[f"- {item['text']}" for item in section["items"]]])
    lines.extend(["\n## 业务风险", *[f"- {item['title']}（{item['level']}）：{'；'.join(item['evidence'])}" for item in risks], "\n## 建议", *[f"- {item['name']}：{item['next_step']}" for item in suggestions], "\n## 数据质量与分析限制", f"- {quality.get('summary', '未提供数据质量摘要。')}", *[f"- {item}" for item in quality.get("limitations", [])], *[f"- {item}" for item in result.get("limitations", [])]])
    markdown = "\n".join(lines) + "\n"
    (reports / "report.md").write_text(markdown, encoding="utf-8")
    (reports / "report.html").write_text(f"<html><meta charset='utf-8'><body><pre>{escape(markdown)}</pre></body></html>", encoding="utf-8")
    document = Document()
    document.add_heading("数据分析决策报告", 0)
    for heading, content in (("分析目标", job["objective"]), ("核心结论", result.get("core_conclusion", "未生成核心结论。")), ("关键数据", "\n".join(f"{item['label']}：{item['value']}" for item in result.get("key_metrics", []))), ("详细分析", "\n".join(item["text"] for section in sections for item in section["items"])), ("业务风险", "\n".join(item["title"] for item in risks)), ("建议", "\n".join(item["name"] for item in suggestions)), ("数据质量与分析限制", quality.get("summary", "未提供数据质量摘要"))):
        document.add_heading(heading, level=1)
        document.add_paragraph(content)
    document.save(reports / "report.docx")
    return [
        {"format": "markdown", "download_url": f"/api/{route}/{job['id']}/files/reports/report.md"},
        {"format": "html", "download_url": f"/api/{route}/{job['id']}/files/reports/report.html"},
        {"format": "docx", "download_url": f"/api/{route}/{job['id']}/files/reports/report.docx"},
    ]

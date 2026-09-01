from __future__ import annotations

from pathlib import Path

from docx import Document

from .charting import build_chart_specs
from .execution import execute_plan
from .planning import CompositeAnalysisPlan, build_plan
from .profiling import profile_file, read_tables
from .validation import validate_result


HIGH_REVIEW_TERMS = ("医疗", "病", "法律", "合同", "金融", "贷款", "化工", "施工安全", "安全事故")


def run(job: dict, directory: Path, source: Path | None = None) -> dict:
    source = source or next(directory.glob("source.*"))
    if job["intake"]["kind"] == "document":
        result = analyse_document(source, job["objective"])
    else:
        result = analyse_structured(source, job["objective"])
    result["notebook_cells"] = notebook_cells(job["intake"]["kind"], job["objective"])
    result["validation_status"] = result.pop("status")
    result["status"] = "succeeded"
    return result


def notebook_cells(kind: str, objective: str) -> list[dict]:
    if kind == "document":
        return [{"language": "python", "title": "文档审阅", "code": "from docx import Document\n\ndocument = Document('source.docx')\nstatements = [p.text.strip() for p in document.paragraphs if p.text.strip()]"}]
    return [{"language": "python", "title": "读取与问题分析", "code": f"objective = {objective!r}\n# 根据数据画像和问题选择允许的受控计算。"}]


def analyse_structured(source: Path, objective: str) -> dict:
    profile = profile_file(source)
    tables = dict(read_tables(source))
    plan = build_plan(profile, objective)
    validated = validate_result(execute_plan(tables, plan), profile)
    result = validated.to_dict()
    plan_details = _plan_details(plan)
    if isinstance(plan, CompositeAnalysisPlan) and result["findings"]:
        result["answer"] = _composite_answer(result["findings"], result["status"])
    result.update(
        {
            "analysis": {
                "kind": "structured_analysis",
                "profile": profile.to_dict(),
                "plan": plan_details,
            },
            "core_conclusion": result["answer"],
            "key_metrics": [
                {
                    "label": finding["kind"],
                    "value": _display_metric_value(finding),
                    "detail": finding["conclusion"],
                }
                for finding in result["findings"]
                if finding["metric_value"] is not None
            ],
            "sections": ([{"title": "分析结论", "items": [{"text": finding["conclusion"]} for finding in result["findings"]]}] if result["findings"] else []),
            "data_quality": {
                "summary": "结论仅使用通过证据校验的受控计算结果。",
                "limitations": result["limitations"],
            },
            "chart_specs": build_chart_specs(result["findings"]),
            "charts": [],
        }
    )
    return result


def _display_metric_value(finding: dict) -> str:
    context = finding.get("context", {})
    if context.get("metric_kind") == "ratio" or context.get("related_metric_kind") == "ratio":
        return f"{float(finding['metric_value']) * 100:.2f}%"
    return str(finding["metric_value"])


def _composite_answer(findings: list[dict], validation_status: str) -> str:
    trends = [finding for finding in findings if finding["kind"] == "trend"]
    trend_items = [_format_composite_trend(finding) for finding in trends]
    anomalies = [
        finding
        for finding in findings
        if finding["kind"] == "anomaly" and isinstance(finding["value"], dict) and finding["value"]
    ]
    anomaly_items = [_format_composite_anomaly(finding) for finding in anomalies]
    anomaly_items.extend(
        f"{finding['context']['metric']}：未发现超过 IQR 阈值的月度异常"
        for finding in findings
        if finding["kind"] == "anomaly" and finding["value"] == []
    )
    if not anomaly_items:
        anomaly_items = ["没有足够证据识别月度异常"]
    reasons = [finding for finding in findings if finding["kind"] in {"reason_comment", "reason_driver", "reason_unavailable"}]
    reason_items = [_format_reason(finding) for finding in reasons]
    if not reason_items:
        reason_items = ["没有与异常月份对应的可验证原因线索"]
    completion = (
        "已完成所请求指标的月度分析"
        if validation_status == "SUCCESS"
        else "部分完成所请求的月度分析，以下仅报告已有可验证结果，未完成项见限制"
    )
    return "结论：" + completion + "。趋势：" + "；".join(trend_items) + "。异常月份：" + "；".join(anomaly_items) + "。原因证据：" + "；".join(reason_items) + "。"


def _format_composite_anomaly(finding: dict) -> str:
    metric = finding["context"]["metric"]
    value = finding["value"]
    if finding["context"].get("metric_kind") == "ratio":
        current = float(value["current_value"]) * 100
        preceding = float(value["preceding_value"]) * 100
        return f"{metric} {value['period']}：{current:.2f}%（上期 {preceding:.2f}%，变化 {current - preceding:+.2f} 个百分点）"
    return f"{metric} {value['period']}：{float(value['current_value']):g}（较上期 {value['change_pct']:+g}%）"


def _format_composite_trend(finding: dict) -> str:
    metric = finding["context"]["metric"]
    points = finding["value"]
    first = float(points[0]["value"])
    last = float(points[-1]["value"])
    if finding["context"].get("metric_kind") == "ratio":
        change = (last - first) * 100
        direction = "上升" if change > 0 else "下降" if change < 0 else "持平"
        return f"{metric}从 {first * 100:.2f}% 到 {last * 100:.2f}%（{direction} {abs(change):.2f} 个百分点）"
    change = float(finding["context"].get("first_to_last_change_pct") or 0)
    direction = "上升" if change > 0 else "下降" if change < 0 else "持平"
    return f"{metric}从 {first:g} 到 {last:g}（{direction} {abs(change):.2f}%）"


def _format_reason(finding: dict) -> str:
    value = finding["value"]
    if finding["kind"] == "reason_comment":
        return f"{value['period']}：可能原因线索，{value['field']}“{value['text']}”（源行 {value['source_row']}，非因果证明）"
    if finding["kind"] == "reason_driver":
        return f"{value['period']}：同期联动线索，{value['field']} {value['comparison_value']:g}→{value['current_value']:g}（{value['change_pct']:+g}%，非因果证明）"
    return f"{value['period']}：可用字段无法确定原因"


def _plan_details(plan: object) -> dict:
    if isinstance(plan, CompositeAnalysisPlan):
        return {
            "kind": "composite",
            "question": plan.question,
            "status": plan.status,
            "table": plan.table,
            "operations": list(plan.operations),
            "time_field": plan.time_field,
            "metrics": [
                {"name": metric.name, "kind": metric.kind, "fields": metric.fields, "formula": metric.formula}
                for metric in plan.metrics
            ],
            "reason_fields": list(plan.reason_fields),
            "driver_fields": list(plan.driver_fields),
            "period_start": plan.period_start,
            "period_end": plan.period_end,
        }
    return {
        "kind": "single",
        "question": plan.question,
        "status": plan.status,
        "table": plan.table,
        "operations": list(plan.operations),
        "fields": plan.fields,
        "aggregation": plan.aggregation,
        "parameters": plan.parameters,
    }


def analyse_document(source: Path, objective: str) -> dict:
    document = Document(source)
    statements = [paragraph.text.strip() for paragraph in document.paragraphs if paragraph.text.strip()]
    evidence = [{"level": "document_statement", "summary": f"文档陈述：{statement}"} for statement in statements[:10]]
    conclusion = statements[0] if statements else "文档没有可提取的段落文本，无法回答分析问题。"
    risk = {"title": "文档主张待核验", "object": "文档陈述", "level": "medium", "evidence": ["DOCX 内容为文本陈述，不等同于测量数据或已验证事实。"], "human_review_required": any(term in objective for term in HIGH_REVIEW_TERMS), "mitigation": "将关键主张映射到原始数据、责任人或审计证据后再决策。"}
    return {"status": "SUCCESS", "answer": conclusion, "findings": [], "analysis": {"kind": "document_review", "statement_count": len(statements)}, "core_conclusion": conclusion, "key_metrics": [{"label": "文档陈述", "value": str(len(statements)), "detail": "条"}], "sections": [{"title": "文档证据", "items": [{"text": item["summary"]} for item in evidence]}], "business_risks": [risk], "data_quality": {"summary": "文档内容不包含可验证测量数据。", "limitations": ["需要原始表格数据才能做统计分析。"]}, "charts": [], "evidence": evidence or [{"level": "document_statement", "summary": "文档没有可提取的段落文本。"}], "risks": [risk], "suggestions": [], "options": [], "limitations": ["文档审阅不验证文本中的数字或事实；需要提供原始数据才能做统计推断。"]}

from datetime import datetime
import hashlib
from io import BytesIO
import json

import pandas as pd
from docx import Document
from fastapi.testclient import TestClient
from openpyxl import Workbook

from studio_api.app import app
from studio_api.answering import analyse_spreadsheet
from studio_api.intake import inspect_file, read_spreadsheet
from studio_api.questioning import plan_question


client = TestClient(app)


def test_generic_engine_returns_a_traceable_direct_ranking_answer(tmp_path):
    source = tmp_path / "measurements.csv"
    source.write_text(
        "bucket,measure_x\n"
        "A,12\n"
        "B,20\n"
        "A,15\n",
        encoding="utf-8",
    )

    from studio_api.engine import run

    result = run(
        {"objective": "哪个 bucket 的 measure_x 最高？", "intake": {"kind": "spreadsheet"}},
        tmp_path,
        source,
    )

    assert result["status"] == "succeeded"
    assert result["validation_status"] == "SUCCESS"
    assert "A" in result["answer"]
    assert "27" in result["answer"]
    assert result["findings"][0]["metric_value"] == 27.0
    assert result["findings"][0]["evidence"]["output_value"] == {"bucket": "A", "aggregate": 27.0}
    assert result["findings"][0]["evidence"]["calculation"] == "groupby(bucket).sum(measure_x)"
    assert result["evidence"][0]["source"]["file_hash"]


def test_structured_trend_returns_one_chart_spec_per_requested_metric(tmp_path):
    from studio_api.engine import run

    source = write_multi_metric_financial_workbook(tmp_path)
    result = run(
        {
            "objective": "分析 2024-01 到 2025-12 的营业收入、毛利率和营业利润趋势",
            "intake": {"kind": "spreadsheet"},
        },
        tmp_path,
        source,
    )

    specs = result["chart_specs"]
    assert [spec["series"][0]["name"] for spec in specs] == ["营业收入", "毛利率", "营业利润"]
    assert all(spec["type"] == "line" for spec in specs)
    assert all(spec["series"][0]["points"] for spec in specs)


def test_upload_intake_summarizes_the_same_largest_data_table_as_analysis(tmp_path):
    workbook = Workbook()
    overview = workbook.active
    overview.title = "财务概览"
    overview.append(["项目", "数值"])
    overview.append(["营业收入", 120])
    detail = workbook.create_sheet("经营明细")
    detail.append(["月份", "区域", "营业收入"])
    for index in range(12):
        detail.append([f"2025-{index + 1:02d}", "华东", 100 + index])
    source = tmp_path / "multi-sheet.xlsx"
    workbook.save(source)

    intake = inspect_file(source)

    assert intake["rows"] == 12
    assert intake["columns"] == ["月份", "区域", "营业收入"]


def test_short_chinese_metric_trend_uses_monthly_composite_analysis(tmp_path):
    from studio_api.engine import run

    source = write_multi_metric_financial_workbook(tmp_path)
    result = run({"objective": "分析营业收入趋势", "intake": {"kind": "spreadsheet"}}, tmp_path, source)

    assert result["validation_status"] == "SUCCESS"
    assert result["analysis"]["plan"]["kind"] == "composite"
    assert result["findings"][0]["kind"] == "trend"
    assert "。。" not in result["answer"]


def test_group_comparison_returns_a_complete_category_bar_spec():
    from studio_api.charting import build_chart_specs

    specs = build_chart_specs(
        [
            {
                "kind": "group_comparison",
                "value": [{"group": "华东", "value": 120}, {"group": "华南", "value": 90}],
                "metric_value": 120,
                "context": {"metric": "营业收入"},
                "evidence": {"fields": ["区域", "营业收入"]},
            }
        ]
    )

    assert specs[0]["type"] == "bar"
    assert specs[0]["series"][0]["points"] == [{"x": "华东", "y": 120.0}, {"x": "华南", "y": 90.0}]


def test_validator_rejects_a_numeric_finding_without_calculation_evidence(tmp_path):
    source = tmp_path / "measurements.csv"
    source.write_text("bucket,measure_x\nA,12\nB,20\n", encoding="utf-8")

    from studio_api.execution import ComputedFinding, ExecutionResult, FindingEvidence
    from studio_api.profiling import profile_file
    from studio_api.validation import validate_result

    invalid = ComputedFinding(
        kind="ranking",
        value="B",
        metric_value=20.0,
        conclusion="B 的 measure_x 为 20。",
        confidence=0.8,
        evidence=FindingEvidence(
            source={"file_hash": profile_file(source).file_hash, "table": "measurements"},
            fields=("bucket", "measure_x"),
            filters=(),
            grouping=("bucket",),
            calculation="",
            row_indices=(1,),
        ),
    )

    result = validate_result(ExecutionResult("SUCCESS", (invalid,), ()), profile_file(source))

    assert result.status == "INSUFFICIENT_DATA"
    assert result.findings == ()
    assert any("calculation" in limitation for limitation in result.limitations)


def test_generic_engine_explains_missing_required_fields_without_fixed_defaults(tmp_path):
    source = tmp_path / "measurements.csv"
    source.write_text("bucket,measure_x\nA,12\nB,20\n", encoding="utf-8")

    from studio_api.engine import run

    result = run(
        {"objective": "分析 measure_x 趋势", "intake": {"kind": "spreadsheet"}},
        tmp_path,
        source,
    )

    assert result["status"] == "succeeded"
    assert result["validation_status"] == "INSUFFICIENT_DATA"
    assert "time" in result["answer"]
    assert result["findings"] == []


def test_profile_classifies_nonstandard_order_fields_without_fixed_business_names(tmp_path):
    source = tmp_path / "orders.csv"
    source.write_text(
        "dt,col1,value_x\n"
        "2025-01-01,alpha,10\n"
        "2025-01-02,beta,12\n"
        "2025-01-03,alpha,11\n",
        encoding="utf-8",
    )

    from studio_api.profiling import profile_file

    profile = profile_file(source)
    columns = {column.name: column for column in profile.tables[0].columns}

    assert profile.file_hash == hashlib.sha256(source.read_bytes()).hexdigest()
    assert columns["dt"].semantic_role == "time"
    assert columns["dt"].confidence >= 0.70
    assert columns["col1"].semantic_role == "dimension"
    assert columns["col1"].confidence >= 0.70
    assert columns["value_x"].semantic_role == "metric"
    assert columns["value_x"].confidence >= 0.70


def test_profile_serializes_nested_json_values_without_failing_unique_measurement(tmp_path):
    source = tmp_path / "events.json"
    source.write_text(
        json.dumps([{"dt": "2025-01-01", "tags": ["north", "urgent"], "value_x": 10}]),
        encoding="utf-8",
    )

    from studio_api.profiling import profile_file

    profile = profile_file(source)
    tags = next(column for column in profile.tables[0].columns if column.name == "tags")

    assert tags.semantic_role == "uncertain"
    assert tags.confidence < 0.70
    assert tags.samples == ['["north","urgent"]']
    assert tags.unique_ratio == 1.0


def test_relationships_auto_use_matching_unique_ids_across_excel_sheets(tmp_path):
    source = tmp_path / "customers.xlsx"
    workbook = Workbook()
    customers = workbook.active
    customers.title = "customers"
    customers.append(["customer_id", "name"])
    customers.append([101, "A"])
    customers.append([102, "B"])
    orders = workbook.create_sheet("orders")
    orders.append(["customer_id", "order_value"])
    orders.append([101, 10])
    orders.append([102, 20])
    workbook.save(source)

    from studio_api.profiling import profile_file
    from studio_api.relationships import discover_relationships

    candidates = discover_relationships([profile_file(source)])

    assert len(candidates) == 1
    candidate = candidates[0]
    assert candidate.left_table == "customers"
    assert candidate.right_table == "orders"
    assert candidate.left_field == candidate.right_field == "customer_id"
    assert candidate.relation_type == "key_match"
    assert candidate.confidence >= 0.85
    assert candidate.can_auto_use is True
    assert candidate.requires_confirmation is False


def test_relationships_do_not_auto_use_similar_low_quality_keys(tmp_path):
    source = tmp_path / "ambiguous.xlsx"
    workbook = Workbook()
    left = workbook.active
    left.title = "left"
    left.append(["customer_id"])
    left.append([101])
    left.append([101])
    left.append([102])
    right = workbook.create_sheet("right")
    right.append(["customer-id"])
    right.append([101])
    right.append([103])
    right.append([103])
    workbook.save(source)

    from studio_api.profiling import profile_file
    from studio_api.relationships import discover_relationships

    candidates = discover_relationships([profile_file(source)])

    assert len(candidates) == 1
    candidate = candidates[0]
    assert candidate.confidence < 0.85
    assert candidate.can_auto_use is False
    assert candidate.requires_confirmation is True


def test_batch_profiles_find_matching_unique_ids_in_independent_files(tmp_path):
    customers = tmp_path / "customers.csv"
    orders = tmp_path / "orders.csv"
    customers.write_text("customer_id,name\n101,A\n102,B\n", encoding="utf-8")
    orders.write_text("customer_id,order_value\n101,10\n102,20\n", encoding="utf-8")

    from studio_api.profiling import profile_files
    from studio_api.relationships import discover_relationships

    profiles = profile_files([customers, orders])
    candidates = discover_relationships(profiles)

    assert len(profiles) == 2
    assert {profile.file_hash for profile in profiles} == {
        hashlib.sha256(customers.read_bytes()).hexdigest(),
        hashlib.sha256(orders.read_bytes()).hexdigest(),
    }
    assert len(candidates) == 1
    assert candidates[0].can_auto_use is True
    assert candidates[0].left_source != candidates[0].right_source


def test_relationships_return_empty_when_no_normalized_name_or_value_overlap(tmp_path):
    first = tmp_path / "first.csv"
    second = tmp_path / "second.csv"
    first.write_text("account_id\n101\n102\n", encoding="utf-8")
    second.write_text("device_code\n301\n302\n", encoding="utf-8")

    from studio_api.profiling import profile_files
    from studio_api.relationships import discover_relationships

    assert discover_relationships(profile_files([first, second])) == []


def test_relationships_do_not_merge_different_explicit_identifier_names(tmp_path):
    accounts = tmp_path / "accounts.csv"
    groups = tmp_path / "groups.csv"
    accounts.write_text("account_id\n101\n102\n", encoding="utf-8")
    groups.write_text("group_id\n101\n102\n", encoding="utf-8")

    from studio_api.profiling import profile_files
    from studio_api.relationships import discover_relationships

    assert discover_relationships(profile_files([accounts, groups])) == []


def test_non_identifier_field_with_id_substring_requires_confirmation(tmp_path):
    first = tmp_path / "first.csv"
    second = tmp_path / "second.csv"
    first.write_text("paid_amount\n10\n20\n", encoding="utf-8")
    second.write_text("paid_amount\n10\n20\n", encoding="utf-8")

    from studio_api.profiling import profile_files
    from studio_api.relationships import discover_relationships

    profiles = profile_files([first, second])
    field = profiles[0].tables[0].columns[0]
    candidates = discover_relationships(profiles)

    assert field.semantic_role == "metric"
    assert len(candidates) == 1
    assert candidates[0].can_auto_use is False
    assert candidates[0].requires_confirmation is True


def test_plan_ranks_the_requested_dimension_and_metric(tmp_path):
    source = tmp_path / "orders.csv"
    frame = pd.DataFrame(
        {
            "product_name": ["A", "B", "A"],
            "sales_amount": [12, 20, 15],
        }
    )
    frame.to_csv(source, index=False)

    from studio_api.execution import execute_plan
    from studio_api.planning import build_plan
    from studio_api.profiling import profile_file

    profile = profile_file(source)
    plan = build_plan(profile, "哪个产品销售额最高？")
    result = execute_plan({profile.tables[0].name: frame}, plan)

    assert plan.status == "READY"
    assert plan.operations == ("ranking",)
    assert plan.fields == {"dimension": "product_name", "metric": "sales_amount"}
    assert result.status == "SUCCESS"
    assert result.findings[0].value == "A"
    assert result.findings[0].metric_value == 27.0
    assert result.findings[0].evidence.calculation == "groupby(product_name).sum(sales_amount)"
    assert result.findings[0].evidence.source == {
        "file_hash": profile.file_hash,
        "table": profile.tables[0].name,
    }


def test_plan_matches_high_confidence_chinese_fields_from_the_question(tmp_path):
    source = tmp_path / "orders-cn.csv"
    frame = pd.DataFrame(
        {
            "产品": ["A", "B", "A"],
            "销售额": [12, 20, 15],
        }
    )
    frame.to_csv(source, index=False)

    from studio_api.planning import build_plan
    from studio_api.profiling import profile_file

    plan = build_plan(profile_file(source), "哪个产品销售额最高？")

    assert plan.status == "READY"
    assert plan.operations == ("ranking",)
    assert plan.fields == {"dimension": "产品", "metric": "销售额"}


def test_plan_builds_composite_chinese_metric_analysis_from_one_profiled_table(tmp_path):
    source = tmp_path / "monthly-finance.csv"
    frame = pd.DataFrame(
        {
            "期间": ["2025-01", "2025-02", "2025-03"],
            "营业收入（万元）": [100, 120, 90],
            "毛利（万元）": [30, 36, 18],
            "营业利润（万元）": [15, 18, 6],
            "备注": ["正常", "正常", "需求回落"],
        }
    )
    frame.to_csv(source, index=False)

    from studio_api.planning import CompositeAnalysisPlan, build_plan
    from studio_api.profiling import profile_file

    plan = build_plan(
        profile_file(source),
        "按期间分析营业收入、毛利率和营业利润趋势，指出异常期间及可能原因，并引用备注证据。",
    )

    assert isinstance(plan, CompositeAnalysisPlan)
    assert plan.table == "monthly-finance"
    assert plan.time_field == "期间"
    assert plan.operations == ("trend", "anomaly", "reason_evidence")
    assert len(plan.metrics) == 3
    assert not hasattr(plan, "metric")
    assert [(metric.name, metric.kind, metric.fields) for metric in plan.metrics] == [
        ("营业收入", "direct", {"metric": "营业收入（万元）"}),
        ("毛利率", "ratio", {"numerator": "毛利（万元）", "denominator": "营业收入（万元）"}),
        ("营业利润", "direct", {"metric": "营业利润（万元）"}),
    ]
    assert plan.metrics[1].formula == "sum(毛利（万元）) / sum(营业收入（万元）)"


def test_composite_plan_does_not_use_a_margin_column_as_gross_profit_amount(tmp_path):
    source = tmp_path / "margin-only.csv"
    pd.DataFrame(
        {
            "期间": ["2025-01", "2025-02", "2025-03"],
            "营业收入": [100, 120, 90],
            "毛利率": [0.30, 0.30, 0.20],
            "营业利润": [15, 18, 6],
        }
    ).to_csv(source, index=False)

    from studio_api.planning import build_plan
    from studio_api.profiling import profile_file

    plan = build_plan(profile_file(source), "按期间分析营业收入、毛利率和营业利润趋势，识别异常。")
    gross_margin = next(metric for metric in plan.metrics if metric.name == "毛利率")

    assert gross_margin.fields == {"denominator": "营业收入"}
    assert gross_margin.missing_fields == ("numerator",)
    assert gross_margin.formula is None


def test_composite_plan_does_not_treat_growth_or_margin_rates_as_direct_amounts(tmp_path):
    source = tmp_path / "rates-only.csv"
    pd.DataFrame(
        {
            "期间": ["2025-01", "2025-02", "2025-03"],
            "营业收入增长率": [0.10, 0.12, -0.05],
            "毛利": [30, 36, 18],
            "营业利润率": [0.15, 0.15, 0.07],
        }
    ).to_csv(source, index=False)

    from studio_api.planning import build_plan
    from studio_api.profiling import profile_file

    plan = build_plan(profile_file(source), "按期间分析营业收入、毛利率和营业利润趋势，识别异常。")

    assert next(metric for metric in plan.metrics if metric.name == "营业收入").missing_fields == ("metric",)
    assert next(metric for metric in plan.metrics if metric.name == "毛利率").missing_fields == ("denominator",)
    assert next(metric for metric in plan.metrics if metric.name == "营业利润").missing_fields == ("metric",)


def test_composite_plan_does_not_treat_a_share_ratio_as_revenue_amount(tmp_path):
    source = tmp_path / "share-ratio-only.csv"
    pd.DataFrame(
        {
            "期间": ["2025-01", "2025-02", "2025-03"],
            "营业收入占比": [0.40, 0.50, 0.30],
            "毛利": [30, 36, 18],
            "营业利润": [15, 18, 6],
        }
    ).to_csv(source, index=False)

    from studio_api.planning import build_plan
    from studio_api.profiling import profile_file

    plan = build_plan(profile_file(source), "按期间分析营业收入、毛利率和营业利润趋势，识别异常。")

    assert next(metric for metric in plan.metrics if metric.name == "营业收入").missing_fields == ("metric",)
    gross_margin = next(metric for metric in plan.metrics if metric.name == "毛利率")
    assert gross_margin.fields == {"numerator": "毛利"}
    assert gross_margin.missing_fields == ("denominator",)
    assert gross_margin.formula is None


def test_composite_plan_uses_a_partial_table_with_time_over_a_complete_table_without_time(tmp_path):
    source = tmp_path / "two-tables.json"
    source.write_text(
        json.dumps(
            {
                "complete_without_time": [
                    {"营业收入": 100, "毛利": 30, "营业利润": 15},
                    {"营业收入": 120, "毛利": 36, "营业利润": 18},
                ],
                "partial_with_time": [
                    {"期间": "2025-01", "营业收入": 100, "营业利润": 15},
                    {"期间": "2025-02", "营业收入": 120, "营业利润": 18},
                ],
            },
            ensure_ascii=False,
        ),
        encoding="utf-8",
    )

    from studio_api.planning import build_plan
    from studio_api.profiling import profile_file

    plan = build_plan(profile_file(source), "按期间分析营业收入、毛利率和营业利润趋势，识别异常。")

    assert plan.table == "partial_with_time"
    assert plan.time_field == "期间"
    assert plan.status == "PARTIAL"


def test_composite_plan_keeps_only_high_confidence_reason_evidence_fields_and_reports_missing_notes(tmp_path):
    source = tmp_path / "financial-notes.csv"
    no_notes_source = tmp_path / "financial-no-notes.csv"
    pd.DataFrame(
        {
            "期间": ["2025-01", "2025-02", "2025-03"],
            "营业收入": [100, 120, 90],
            "毛利": [30, 36, 18],
            "营业利润": [15, 18, 6],
            "备注": ["正常", "正常", "需求回落"],
            "说明": ["仅一条", None, None],
        }
    ).to_csv(source, index=False)
    pd.DataFrame(
        {
            "期间": ["2025-01", "2025-02", "2025-03"],
            "营业收入": [100, 120, 90],
            "毛利": [30, 36, 18],
            "营业利润": [15, 18, 6],
        }
    ).to_csv(no_notes_source, index=False)

    from studio_api.planning import build_plan
    from studio_api.profiling import profile_file

    profile_with_notes = profile_file(source)
    assert next(column for column in profile_with_notes.tables[0].columns if column.name == "说明").confidence < 0.70

    with_notes = build_plan(profile_with_notes, "按期间分析营业收入、毛利率和营业利润趋势，识别异常并引用备注证据。")
    without_notes = build_plan(
        profile_file(no_notes_source), "按期间分析营业收入、毛利率和营业利润趋势，识别异常并引用原因证据。"
    )

    assert with_notes.reason_fields == ("备注",)
    assert without_notes.reason_fields == ()
    assert any("备注/说明/原因证据字段" in limitation for limitation in without_notes.limitations)


def test_composite_planning_uses_monthly_analysis_for_explicit_metric_trends(tmp_path):
    source = tmp_path / "financial.csv"
    pd.DataFrame(
        {
            "期间": ["2025-01", "2025-02", "2025-03"],
            "营业收入": [100, 120, 90],
            "毛利": [30, 36, 18],
            "营业利润": [15, 18, 6],
        }
    ).to_csv(source, index=False)

    from studio_api.planning import CompositeAnalysisPlan, build_plan
    from studio_api.profiling import profile_file

    profile = profile_file(source)
    no_explicit_time = build_plan(profile, "分析营业收入、营业利润趋势并识别异常。")
    simple = build_plan(profile, "分析营业收入趋势")

    assert isinstance(no_explicit_time, CompositeAnalysisPlan)
    assert isinstance(simple, CompositeAnalysisPlan)


def test_composite_plan_supports_one_explicit_generic_metric_with_time_range_and_evidence_request(tmp_path):
    source = tmp_path / "toll.csv"
    pd.DataFrame(
        {
            "data_month": pd.period_range("2022-01", "2023-01", freq="M").astype(str),
            "opma_section_id": ["G03213717"] * 13,
            "total_fee": [100, 102, 101, 103, 105, 104, 106, 300, 107, 108, 109, 110, 111],
        }
    ).to_csv(source, index=False)

    from studio_api.planning import CompositeAnalysisPlan, build_plan
    from studio_api.profiling import profile_file

    plan = build_plan(
        profile_file(source),
        "分析 2022年1月到2023年1月各路段总通行费用（total_fee）的趋势，指出异常月份及可能原因，并引用数据证据。",
    )

    assert isinstance(plan, CompositeAnalysisPlan)
    assert plan.time_field == "data_month"
    assert plan.operations == ("trend", "anomaly", "reason_evidence")
    assert [(metric.name, metric.fields) for metric in plan.metrics] == [("total_fee", {"metric": "total_fee"})]
    assert (plan.period_start, plan.period_end) == ("2022-01", "2023-01")


def test_composite_execution_aggregates_monthly_metrics_and_ratio_before_trend_calculation(tmp_path):
    periods = pd.period_range("2024-01", "2025-12", freq="M").astype(str)
    rows = []
    for index, period in enumerate(periods):
        revenue = 100 + index * 5
        gross_profit = 45 + index * 2
        operating_profit = 20 + index
        rows.extend(
            [
                {"期间": period, "营业收入": revenue - 10, "毛利": gross_profit, "营业利润": operating_profit - 2},
                {"期间": period, "营业收入": 10, "毛利": 0, "营业利润": 2},
            ]
        )
    frame = pd.DataFrame(rows)
    source = tmp_path / "monthly-finance.csv"
    frame.to_csv(source, index=False)

    from studio_api.execution import execute_plan
    from studio_api.planning import build_plan
    from studio_api.profiling import profile_file

    profile = profile_file(source)
    plan = build_plan(profile, "按期间分析营业收入、毛利率和营业利润趋势，识别异常。")
    result = execute_plan({profile.tables[0].name: frame}, plan)

    revenue_trend = next(finding for finding in result.findings if finding.kind == "trend" and finding.context["metric"] == "营业收入")
    margin_trend = next(finding for finding in result.findings if finding.kind == "trend" and finding.context["metric"] == "毛利率")

    assert result.status == "SUCCESS"
    assert revenue_trend.value[0] == {"period": "2024-01", "value": 100.0}
    assert margin_trend.value[0] == {"period": "2024-01", "value": 0.45}
    assert margin_trend.value[0]["value"] != 0.225
    assert revenue_trend.context["first_to_last_change_pct"] == 115.0
    assert revenue_trend.context["maximum"] == 215.0
    assert revenue_trend.context["minimum"] == 100.0
    assert revenue_trend.evidence.calculation == "monthly_sum(营业收入)"
    assert margin_trend.evidence.calculation == "monthly_sum(毛利) / monthly_sum(营业收入)"
    assert revenue_trend.evidence.fields == ("期间", "营业收入")
    assert revenue_trend.evidence.grouping == ("期间",)
    assert revenue_trend.evidence.row_indices == tuple(frame.index)


def test_composite_execution_detects_iqr_anomaly_from_monthly_aggregate_changes(tmp_path):
    periods = pd.period_range("2024-01", "2025-12", freq="M").astype(str)
    rows = []
    for index, period in enumerate(periods):
        revenue = 100 + index * 5
        operating_profit = 20 + index
        if period == "2025-07":
            revenue = 600
            operating_profit = 120
        rows.extend(
            [
                {"期间": period, "营业收入": revenue - 10, "毛利": (revenue - 10) * 0.4, "营业利润": operating_profit - 2},
                {"期间": period, "营业收入": 10, "毛利": 0, "营业利润": 2},
            ]
        )
    frame = pd.DataFrame(rows)
    source = tmp_path / "monthly-anomaly.csv"
    frame.to_csv(source, index=False)

    from studio_api.execution import execute_plan
    from studio_api.planning import build_plan
    from studio_api.profiling import profile_file

    profile = profile_file(source)
    plan = build_plan(profile, "按期间分析营业收入、毛利率和营业利润趋势，识别异常。")
    result = execute_plan({profile.tables[0].name: frame}, plan)

    anomaly = next(
        finding
        for finding in result.findings
        if finding.kind == "anomaly"
        and finding.context["metric"] == "营业收入"
        and finding.context["period"] == "2025-07"
    )

    assert anomaly.metric_value == 600.0
    assert anomaly.value == {"period": "2025-07", "current_value": 600.0, "preceding_value": 185.0, "change_pct": 224.3243}
    assert anomaly.context["method"] == "IQR"
    assert anomaly.context["threshold"]["multiplier"] == 1.5
    assert anomaly.evidence.calculation == "iqr_outliers(monthly_percent_change(营业收入), k=1.5)"
    assert anomaly.evidence.fields == ("期间", "营业收入")
    assert anomaly.evidence.grouping == ("期间",)

    operating_profit_anomaly = next(
        finding
        for finding in result.findings
        if finding.kind == "anomaly"
        and finding.context["metric"] == "营业利润"
        and finding.context["period"] == "2025-07"
    )

    assert operating_profit_anomaly.value == {
        "period": "2025-07",
        "current_value": 120.0,
        "preceding_value": 37.0,
        "change_pct": 224.3243,
    }
    assert operating_profit_anomaly.context["method"] == "IQR"
    assert operating_profit_anomaly.context["threshold"]["multiplier"] == 1.5
    assert operating_profit_anomaly.evidence.calculation == "iqr_outliers(monthly_percent_change(营业利润), k=1.5)"
    assert operating_profit_anomaly.evidence.fields == ("期间", "营业利润")
    assert operating_profit_anomaly.evidence.grouping == ("期间",)


def test_composite_execution_reports_insufficient_monthly_evidence_for_short_anomaly_series(tmp_path):
    frame = pd.DataFrame(
        {
            "期间": ["2025-01", "2025-01", "2025-02", "2025-02", "2025-03", "2025-03", "2025-04", "2025-04"],
            "营业收入": [90, 10, 100, 10, 110, 10, 120, 10],
            "毛利": [36, 0, 40, 0, 44, 0, 48, 0],
            "营业利润": [18, 2, 20, 2, 22, 2, 24, 2],
        }
    )
    source = tmp_path / "short-monthly-finance.csv"
    frame.to_csv(source, index=False)

    from studio_api.execution import execute_plan
    from studio_api.planning import build_plan
    from studio_api.profiling import profile_file

    profile = profile_file(source)
    plan = build_plan(profile, "按期间分析营业收入、毛利率和营业利润趋势，识别异常。")
    result = execute_plan({profile.tables[0].name: frame}, plan)

    assert result.status == "PARTIAL"
    assert all(finding.kind == "trend" for finding in result.findings)
    assert any("有效相邻变化不足四个" in limitation and "零基期" in limitation for limitation in result.limitations)


def test_composite_engine_uses_planned_remark_as_same_period_reason_clue(tmp_path):
    source = tmp_path / "remarked-finance.csv"
    pd.DataFrame(
        {
            "期间": pd.period_range("2025-01", "2025-06", freq="M").astype(str),
            "营业收入": [100, 101, 102, 103, 104, 240],
            "营业利润": [15, 15.1, 15.2, 15.3, 15.4, 36],
            "备注": ["正常经营", "正常经营", "正常经营", "正常经营", "正常经营", "临时大额订单确认"],
        }
    ).to_csv(source, index=False)

    from studio_api.engine import run

    result = run(
        {"objective": "按期间分析营业收入和营业利润趋势，指出异常月份及可能原因", "intake": {"kind": "spreadsheet"}},
        tmp_path,
        source,
    )

    clue = next(item for item in result["findings"] if item["kind"] == "reason_comment")
    assert clue["value"]["period"] == "2025-06"
    assert clue["value"]["text"] == "临时大额订单确认"
    assert clue["value"]["source_row"] == 5
    assert clue["evidence"]["row_indices"] == [5]
    assert "可能原因线索" in result["answer"]
    assert "因果证明" in result["answer"]


def test_composite_engine_reports_same_period_driver_co_movement_with_values_and_source(tmp_path):
    source = tmp_path / "driver-finance.csv"
    pd.DataFrame(
        {
            "期间": pd.period_range("2025-01", "2025-06", freq="M").astype(str),
            "营业收入": [100, 101, 102, 103, 104, 240],
            "营业利润": [15, 15.1, 15.2, 15.3, 15.4, 36],
            "销量": [10, 10, 10, 10, 10, 24],
        }
    ).to_csv(source, index=False)

    from studio_api.engine import run

    result = run(
        {"objective": "按期间分析营业收入和营业利润趋势，指出异常月份及可能原因", "intake": {"kind": "spreadsheet"}},
        tmp_path,
        source,
    )

    clue = next(item for item in result["findings"] if item["kind"] == "reason_driver")
    assert clue["value"] == {
        "period": "2025-06",
        "field": "销量",
        "comparison_value": 10.0,
        "current_value": 24.0,
        "change_pct": 140.0,
        "source_rows": [4, 5],
    }
    assert clue["evidence"]["calculation"] == "monthly_sum(销量); monthly_percent_change(销量)"
    assert clue["evidence"]["row_indices"] == [4, 5]
    assert "同期联动线索" in result["answer"]
    assert "因果证明" in result["answer"]


def test_composite_engine_states_fields_cannot_determine_cause_without_reason_evidence(tmp_path):
    source = tmp_path / "no-reason-finance.csv"
    pd.DataFrame(
        {
            "期间": pd.period_range("2025-01", "2025-06", freq="M").astype(str),
            "营业收入": [100, 101, 102, 103, 104, 240],
            "营业利润": [15, 15.1, 15.2, 15.3, 15.4, 36],
        }
    ).to_csv(source, index=False)

    from studio_api.engine import run

    result = run(
        {"objective": "按期间分析营业收入和营业利润趋势，指出异常月份及可能原因", "intake": {"kind": "spreadsheet"}},
        tmp_path,
        source,
    )

    clue = next(item for item in result["findings"] if item["kind"] == "reason_unavailable")
    assert clue["value"]["period"] == "2025-06"
    assert clue["evidence"]["row_indices"] == [4, 5]
    assert "可用字段无法确定原因" in result["answer"]
    assert "导致" not in result["answer"]


def test_composite_plan_and_execution_support_any_explicit_high_confidence_numeric_metrics(tmp_path):
    source = tmp_path / "operations.csv"
    frame = pd.DataFrame(
        {
            "月份": pd.period_range("2025-01", "2025-08", freq="M").astype(str),
            "销量": [100, 102, 104, 106, 108, 300, 112, 114],
            "成本": [50, 51, 52, 53, 54, 150, 56, 57],
            "退货金额": [5, 5.2, 5.1, 5.3, 5.4, 30, 5.6, 5.7],
        }
    )
    frame.to_csv(source, index=False)

    from studio_api.execution import execute_plan
    from studio_api.planning import CompositeAnalysisPlan, build_plan
    from studio_api.profiling import profile_file

    profile = profile_file(source)
    plan = build_plan(profile, "按月份分析销量、成本和退货金额趋势，指出异常月份及可能原因。")
    result = execute_plan({profile.tables[0].name: frame}, plan)

    assert isinstance(plan, CompositeAnalysisPlan)
    assert [(metric.name, metric.kind, metric.fields) for metric in plan.metrics] == [
        ("销量", "direct", {"metric": "销量"}),
        ("成本", "direct", {"metric": "成本"}),
        ("退货金额", "direct", {"metric": "退货金额"}),
    ]
    assert {finding.context["metric"] for finding in result.findings if finding.kind == "trend"} == {
        "销量",
        "成本",
        "退货金额",
    }
    assert all(
        finding.evidence.calculation == f"monthly_sum({finding.context['metric']})"
        for finding in result.findings
        if finding.kind == "trend"
    )


def test_composite_explicit_metric_matching_prefers_the_longest_field_at_the_same_position(tmp_path):
    source = tmp_path / "prefixed-metrics.csv"
    pd.DataFrame(
        {
            "月份": pd.period_range("2025-01", "2025-06", freq="M").astype(str),
            "销量目标": [100, 105, 110, 115, 120, 125],
            "销量": [90, 95, 100, 105, 110, 115],
            "成本": [50, 52, 54, 56, 58, 60],
        }
    ).to_csv(source, index=False)

    from studio_api.planning import CompositeAnalysisPlan, build_plan
    from studio_api.profiling import profile_file

    plan = build_plan(profile_file(source), "按月份分析销量目标和成本趋势。")

    assert isinstance(plan, CompositeAnalysisPlan)
    assert [metric.name for metric in plan.metrics] == ["销量目标", "成本"]


def test_composite_metric_matching_prefers_a_longer_generic_field_over_a_static_financial_alias(tmp_path):
    source = tmp_path / "financial-prefixed-metrics.csv"
    pd.DataFrame(
        {
            "月份": pd.period_range("2025-01", "2025-06", freq="M").astype(str),
            "营业收入目标": [100, 105, 110, 115, 120, 125],
            "营业收入": [90, 95, 100, 105, 110, 115],
            "成本": [50, 52, 54, 56, 58, 60],
        }
    ).to_csv(source, index=False)

    from studio_api.planning import CompositeAnalysisPlan, build_plan
    from studio_api.profiling import profile_file

    plan = build_plan(profile_file(source), "按月份分析营业收入目标和成本趋势。")

    assert isinstance(plan, CompositeAnalysisPlan)
    assert [(metric.name, metric.fields) for metric in plan.metrics] == [
        ("营业收入目标", {"metric": "营业收入目标"}),
        ("成本", {"metric": "成本"}),
    ]


def test_composite_reason_keeps_all_unique_comments_and_a_significant_numeric_clue(tmp_path):
    periods = list(pd.period_range("2025-01", "2025-07", freq="M").astype(str))
    rows = [
        {"期间": period, "营业收入": 100 + index, "营业利润": 20 + index, "销量": 8 + index, "备注": "正常"}
        for index, period in enumerate(periods)
    ]
    rows.extend(
        [
            {"期间": "2025-08", "营业收入": 100, "营业利润": 30, "销量": 14, "备注": "促销活动"},
            {"期间": "2025-08", "营业收入": 100, "营业利润": 30, "销量": 14, "备注": "供应调整"},
            {"期间": "2025-08", "营业收入": 100, "营业利润": 30, "销量": 14, "备注": "促销活动"},
        ]
    )
    source = tmp_path / "multi-clue.csv"
    pd.DataFrame(rows).to_csv(source, index=False)

    from studio_api.engine import run

    result = run(
        {"objective": "按期间分析营业收入和营业利润趋势，指出异常月份及可能原因", "intake": {"kind": "spreadsheet"}},
        tmp_path,
        source,
    )

    comments = [finding for finding in result["findings"] if finding["kind"] == "reason_comment"]
    drivers = [finding for finding in result["findings"] if finding["kind"] == "reason_driver"]
    assert {finding["value"]["text"] for finding in comments} == {"促销活动", "供应调整"}
    assert len(comments) == 2
    assert len(drivers) == 1
    assert drivers[0]["value"]["field"] == "销量"
    assert drivers[0]["value"]["change_pct"] == 200.0
    assert "可能原因线索" in result["answer"] and "同期联动线索" in result["answer"]


def test_composite_reason_keeps_every_significant_same_period_driver(tmp_path):
    source = tmp_path / "all-significant-drivers.csv"
    pd.DataFrame(
        {
            "期间": pd.period_range("2025-01", "2025-08", freq="M").astype(str),
            "营业收入": [100, 101, 102, 103, 104, 105, 106, 300],
            "营业利润": [20, 20.2, 20.4, 20.6, 20.8, 21, 21.2, 60],
            "销量": [10, 10, 10, 10, 10, 10, 10, 30],
            "单价": [10, 10, 10, 10, 10, 10, 10, 20],
            "成本": [50, 50, 50, 50, 50, 50, 50, 100],
        }
    ).to_csv(source, index=False)

    from studio_api.engine import run

    result = run(
        {"objective": "按期间分析营业收入和营业利润趋势，指出异常月份及可能原因", "intake": {"kind": "spreadsheet"}},
        tmp_path,
        source,
    )

    drivers = [finding for finding in result["findings"] if finding["kind"] == "reason_driver"]
    assert [finding["value"]["field"] for finding in drivers] == ["销量", "单价", "成本"]
    assert all(finding["evidence"]["row_indices"] == [6, 7] for finding in drivers)
    assert all("非因果证明" in finding["conclusion"] for finding in drivers)


def test_composite_reason_skips_a_zero_baseline_driver_without_failing(tmp_path):
    source = tmp_path / "zero-baseline-driver.csv"
    pd.DataFrame(
        {
            "期间": pd.period_range("2025-01", "2025-08", freq="M").astype(str),
            "营业收入": [100, 101, 102, 103, 104, 105, 106, 300],
            "营业利润": [20, 20.2, 20.4, 20.6, 20.8, 21, 21.2, 60],
            "销量": [10, 10, 10, 10, 10, 10, 0, 30],
        }
    ).to_csv(source, index=False)

    from studio_api.engine import run

    result = run(
        {"objective": "按期间分析营业收入和营业利润趋势，指出异常月份及可能原因", "intake": {"kind": "spreadsheet"}},
        tmp_path,
        source,
    )

    assert result["status"] == "succeeded"
    assert not any(finding["kind"] == "reason_driver" for finding in result["findings"])
    assert any(finding["kind"] == "reason_unavailable" for finding in result["findings"])


def test_composite_time_range_filters_before_aggregation_and_is_recorded_in_evidence(tmp_path):
    source = tmp_path / "ranged-metrics.csv"
    frame = pd.DataFrame(
        {
            "月份": pd.period_range("2025-01", "2025-10", freq="M").astype(str),
            "销量": [900, 800, 100, 101, 102, 103, 104, 105, 700, 600],
            "成本": [450, 400, 50, 51, 52, 53, 54, 55, 350, 300],
        }
    )
    frame.to_csv(source, index=False)

    from studio_api.execution import execute_plan
    from studio_api.planning import build_plan
    from studio_api.profiling import profile_file

    profile = profile_file(source)
    plan = build_plan(profile, "分析 2025-03 至 2025-08 的销量和成本趋势并识别异常。")
    result = execute_plan({profile.tables[0].name: frame}, plan)
    trend = next(finding for finding in result.findings if finding.kind == "trend" and finding.context["metric"] == "销量")

    assert plan.period_start == "2025-03"
    assert plan.period_end == "2025-08"
    assert [point["period"] for point in trend.value] == ["2025-03", "2025-04", "2025-05", "2025-06", "2025-07", "2025-08"]
    assert trend.evidence.filters == ("月份：2025-03 至 2025-08",)
    assert trend.evidence.row_indices == (2, 3, 4, 5, 6, 7)


def test_composite_period_comparison_does_not_invent_a_continuous_range(tmp_path):
    source = tmp_path / "compared-months.csv"
    frame = pd.DataFrame(
        {
            "月份": pd.period_range("2025-01", "2025-10", freq="M").astype(str),
            "销量": list(range(100, 110)),
            "成本": list(range(50, 60)),
        }
    )
    frame.to_csv(source, index=False)

    from studio_api.execution import execute_plan
    from studio_api.planning import build_plan
    from studio_api.profiling import profile_file

    profile = profile_file(source)
    plan = build_plan(profile, "按月份比较 2025-03 和 2025-08 的销量和成本趋势。")
    result = execute_plan({profile.tables[0].name: frame}, plan)
    trend = next(finding for finding in result.findings if finding.kind == "trend" and finding.context["metric"] == "销量")

    assert plan.period_start is None and plan.period_end is None
    assert [point["period"] for point in trend.value] == list(pd.period_range("2025-01", "2025-10", freq="M").astype(str))
    assert trend.evidence.filters == ()


def test_composite_reverse_period_range_is_rejected_without_reordering_or_filtering(tmp_path):
    source = tmp_path / "reverse-range.csv"
    frame = pd.DataFrame(
        {
            "月份": pd.period_range("2025-01", "2025-10", freq="M").astype(str),
            "销量": list(range(100, 110)),
            "成本": list(range(50, 60)),
        }
    )
    frame.to_csv(source, index=False)

    from studio_api.execution import execute_plan
    from studio_api.planning import build_plan
    from studio_api.profiling import profile_file

    profile = profile_file(source)
    plan = build_plan(profile, "分析 2025-08 至 2025-03 的销量和成本趋势。")
    result = execute_plan({profile.tables[0].name: frame}, plan)
    trend = next(finding for finding in result.findings if finding.kind == "trend" and finding.context["metric"] == "销量")

    assert plan.period_start is None and plan.period_end is None
    assert any("起始月份晚于结束月份" in limitation for limitation in plan.limitations)
    assert [point["period"] for point in trend.value] == list(pd.period_range("2025-01", "2025-10", freq="M").astype(str))
    assert trend.evidence.filters == ()


def test_composite_partial_and_insufficient_results_keep_a_succeeded_task_lifecycle(tmp_path):
    from studio_api.engine import run

    partial_source = tmp_path / "partial.csv"
    pd.DataFrame(
        {
            "期间": pd.period_range("2025-01", "2025-06", freq="M").astype(str),
            "营业收入": [100, 101, 102, 103, 104, 105],
            "营业利润": [20, 21, 22, 23, 24, 25],
        }
    ).to_csv(partial_source, index=False)
    partial = run(
        {"objective": "按期间分析营业收入、毛利率和营业利润趋势，识别异常。", "intake": {"kind": "spreadsheet"}},
        tmp_path,
        partial_source,
    )

    insufficient_source = tmp_path / "insufficient.csv"
    pd.DataFrame({"营业收入": [100, 110], "营业利润": [20, 22]}).to_csv(insufficient_source, index=False)
    insufficient = run(
        {"objective": "按月份分析营业收入和营业利润趋势。", "intake": {"kind": "spreadsheet"}},
        tmp_path,
        insufficient_source,
    )

    assert partial["status"] == "succeeded"
    assert partial["validation_status"] == "PARTIAL"
    assert insufficient["status"] == "succeeded"
    assert insufficient["validation_status"] == "INSUFFICIENT_DATA"


def test_composite_partial_answer_is_honest_and_margin_anomaly_uses_percentage_display(tmp_path):
    from studio_api.engine import run

    partial_source = tmp_path / "partial-answer.csv"
    pd.DataFrame(
        {
            "期间": pd.period_range("2025-01", "2025-06", freq="M").astype(str),
            "营业收入": [100, 101, 102, 103, 104, 105],
            "营业利润": [20, 21, 22, 23, 24, 25],
        }
    ).to_csv(partial_source, index=False)
    partial = run(
        {"objective": "按期间分析营业收入、毛利率和营业利润趋势，识别异常。", "intake": {"kind": "spreadsheet"}},
        tmp_path,
        partial_source,
    )
    assert "部分完成" in partial["answer"]
    assert "已完成所请求指标" not in partial["answer"]

    margin_source = tmp_path / "margin-anomaly.csv"
    revenues = [100, 101, 102, 103, 104, 105, 106, 107]
    margin_rates = [0.30, 0.301, 0.302, 0.303, 0.304, 0.305, 0.90, 0.307]
    pd.DataFrame(
        {
            "期间": pd.period_range("2025-01", "2025-08", freq="M").astype(str),
            "营业收入": revenues,
            "毛利": [revenue * rate for revenue, rate in zip(revenues, margin_rates)],
            "营业利润": [20, 21, 22, 23, 24, 25, 26, 27],
        }
    ).to_csv(margin_source, index=False)
    margin = run(
        {"objective": "按期间分析营业收入、毛利率和营业利润趋势，识别异常。", "intake": {"kind": "spreadsheet"}},
        tmp_path,
        margin_source,
    )
    margin_anomaly = next(
        finding
        for finding in margin["findings"]
        if finding["kind"] == "anomaly"
        and finding["context"]["metric"] == "毛利率"
        and finding["value"].get("period") == "2025-07"
    )
    margin_card = next(
        item
        for item in margin["key_metrics"]
        if item["label"] == "anomaly" and item["detail"] == margin_anomaly["conclusion"]
    )

    assert "毛利率 2025-07：90.00%（上期 30.50%，变化 +59.50 个百分点）" in margin["answer"]
    assert margin_anomaly["conclusion"] == "毛利率 在 2025-07 的月度汇总值为 90.00%，上期为 30.50%，变化 +59.50 个百分点。"
    assert margin_card["value"] == "90.00%"
    assert margin_card["detail"] == margin_anomaly["conclusion"]


def test_ratio_anomaly_reason_cards_use_percentage_display_for_every_reason_kind(tmp_path):
    from studio_api.engine import run

    periods = pd.period_range("2025-01", "2025-08", freq="M").astype(str)
    revenues = [100, 101, 102, 103, 104, 105, 106, 107]
    margin_rates = [0.30, 0.301, 0.302, 0.303, 0.304, 0.305, 0.90, 0.307]
    variants = {
        "reason_comment": {"备注": ["正常"] * 6 + ["异常促销", "正常"]},
        "reason_driver": {"销量": [10] * 6 + [30, 30]},
        "reason_unavailable": {},
    }

    for expected_kind, extra_columns in variants.items():
        source = tmp_path / f"ratio-{expected_kind}.csv"
        pd.DataFrame(
            {
                "期间": periods,
                "营业收入": revenues,
                "毛利": [revenue * rate for revenue, rate in zip(revenues, margin_rates)],
                "营业利润": [20, 21, 22, 23, 24, 25, 26, 27],
                **extra_columns,
            }
        ).to_csv(source, index=False)
        result = run(
            {
                "objective": "按期间分析营业收入、毛利率和营业利润趋势，识别异常及可能原因。",
                "intake": {"kind": "spreadsheet"},
            },
            tmp_path,
            source,
        )
        reason = next(
            finding
            for finding in result["findings"]
            if finding["kind"] == expected_kind
            and finding["context"].get("related_metric") == "毛利率"
            and finding["context"].get("period") == "2025-07"
        )
        card = next(
            item
            for item in result["key_metrics"]
            if item["label"] == expected_kind and item["detail"] == reason["conclusion"]
        )

        assert reason["metric_value"] == 0.9
        assert card["value"] == "90.00%"
        assert card["detail"] == reason["conclusion"]
        assert "关联毛利率 90.00%" in card["detail"]
        assert "0.9" not in card["detail"]


def test_validator_rejects_unknown_fields_missing_rows_and_mismatched_ratio_formula(tmp_path):
    source = tmp_path / "validation-source.csv"
    source.write_text("bucket,measure_x,profit,sales\nA,12,3,10\nB,20,4,10\n", encoding="utf-8")

    from studio_api.execution import ComputedFinding, ExecutionResult, FindingEvidence
    from studio_api.profiling import profile_file
    from studio_api.validation import validate_result

    profile = profile_file(source)
    evidence_source = {"file_hash": profile.file_hash, "table": profile.tables[0].name}
    unknown_field = ComputedFinding(
        kind="ranking", value="B", metric_value=20.0, conclusion="错误字段。", confidence=0.8,
        evidence=FindingEvidence(evidence_source, ("bucket", "invented"), (), ("bucket",), "max(invented)", (1,)),
    )
    missing_row = ComputedFinding(
        kind="ranking", value="B", metric_value=20.0, conclusion="错误行。", confidence=0.8,
        evidence=FindingEvidence(evidence_source, ("bucket", "measure_x"), (), ("bucket",), "max(measure_x)", (99,)),
    )
    mismatched_formula = ComputedFinding(
        kind="trend", value=[{"period": "2025-01", "value": 0.3}], metric_value=0.3,
        conclusion="错误公式。", confidence=0.8,
        evidence=FindingEvidence(evidence_source, ("profit", "sales"), (), (), "monthly_sum(profit) / monthly_sum(sales)", (0, 1)),
        context={
            "metric_kind": "ratio",
            "metric_fields": {"numerator": "profit", "denominator": "sales"},
            "formula": "sum(measure_x) / sum(sales)",
        },
    )

    validated = validate_result(ExecutionResult("SUCCESS", (unknown_field, missing_row, mismatched_formula), ()), profile)

    assert validated.status == "INSUFFICIENT_DATA"
    assert validated.findings == ()
    assert any("源表字段" in limitation for limitation in validated.limitations)
    assert any("源数据行" in limitation for limitation in validated.limitations)
    assert any("分子/分母字段" in limitation for limitation in validated.limitations)


def test_ranking_evidence_excludes_rows_rejected_by_numeric_conversion(tmp_path):
    source = tmp_path / "orders-with-invalid-value.csv"
    frame = pd.DataFrame(
        {
            "product_name": ["A", "A", "A", "A", "A", "B", "B", "B", "B", "B"],
            "sales_amount": [12, "invalid", 15, 0, 0, 20, 1, 1, 1, 1],
        }
    )
    frame.to_csv(source, index=False)

    from studio_api.execution import execute_plan
    from studio_api.planning import build_plan
    from studio_api.profiling import profile_file

    profile = profile_file(source)
    plan = build_plan(profile, "哪个产品销售额最高？")
    result = execute_plan({profile.tables[0].name: frame}, plan)

    assert result.status == "SUCCESS"
    assert result.findings[0].value == "A"
    assert result.findings[0].metric_value == 27.0
    assert result.findings[0].evidence.row_indices == (0, 2, 3, 4)


def test_prediction_without_a_time_field_returns_insufficient_data(tmp_path):
    source = tmp_path / "orders.csv"
    frame = pd.DataFrame(
        {
            "product_name": ["A", "B", "A"],
            "sales_amount": [12, 20, 15],
        }
    )
    frame.to_csv(source, index=False)

    from studio_api.execution import execute_plan
    from studio_api.planning import build_plan
    from studio_api.profiling import profile_file

    profile = profile_file(source)
    plan = build_plan(profile, "预测 sales_amount 未来趋势")
    result = execute_plan({profile.tables[0].name: frame}, plan)

    assert plan.status == "INSUFFICIENT_DATA"
    assert plan.operations == ("forecast",)
    assert plan.fields == {"metric": "sales_amount"}
    assert any("时间字段" in limitation for limitation in plan.limitations)
    assert result.status == "INSUFFICIENT_DATA"
    assert result.findings == ()


def test_generic_anomaly_question_returns_the_outlier_with_row_evidence(tmp_path):
    source = tmp_path / "measurements.csv"
    frame = pd.DataFrame(
        {
            "observed_at": pd.date_range("2025-01-01", periods=5).astype(str),
            "measure_x": [10, 11, 10, 12, 100],
        }
    )
    frame.to_csv(source, index=False)

    from studio_api.execution import execute_plan
    from studio_api.planning import build_plan
    from studio_api.profiling import profile_file

    profile = profile_file(source)
    plan = build_plan(profile, "measure_x 有哪些异常值？")
    result = execute_plan({profile.tables[0].name: frame}, plan)

    assert plan.status == "READY"
    assert plan.operations == ("anomaly",)
    assert plan.fields == {"metric": "measure_x", "time": "observed_at"}
    assert result.status == "SUCCESS"
    assert result.findings[0].value == 100.0
    assert result.findings[0].evidence.calculation == "iqr_outliers(measure_x, k=1.5)"
    assert result.findings[0].evidence.row_indices == (4,)


def test_anomaly_without_optional_context_returns_partial_finding(tmp_path):
    source = tmp_path / "measurements.csv"
    frame = pd.DataFrame({"measure_x": [10, 11, 10, 12, 100]})
    frame.to_csv(source, index=False)

    from studio_api.execution import execute_plan
    from studio_api.planning import build_plan
    from studio_api.profiling import profile_file

    profile = profile_file(source)
    plan = build_plan(profile, "measure_x 有哪些异常值？")
    result = execute_plan({profile.tables[0].name: frame}, plan)

    assert plan.status == "PARTIAL"
    assert any("上下文字段" in limitation for limitation in plan.limitations)
    assert result.status == "PARTIAL"
    assert result.findings[0].value == 100.0


def test_low_confidence_fields_are_not_used_for_a_ranking(tmp_path):
    source = tmp_path / "single-order.csv"
    frame = pd.DataFrame({"product_name": ["A"], "sales_amount": [12]})
    frame.to_csv(source, index=False)

    from studio_api.execution import execute_plan
    from studio_api.planning import build_plan
    from studio_api.profiling import profile_file

    profile = profile_file(source)
    plan = build_plan(profile, "哪个产品销售额最高？")
    result = execute_plan({profile.tables[0].name: frame}, plan)

    assert plan.status == "INSUFFICIENT_DATA"
    assert plan.fields == {}
    assert any("置信度" in limitation for limitation in plan.limitations)
    assert result.status == "INSUFFICIENT_DATA"
    assert result.findings == ()


def test_question_terms_select_the_requested_metric_among_profile_candidates(tmp_path):
    source = tmp_path / "orders.csv"
    frame = pd.DataFrame(
        {
            "product_name": ["A", "B", "A"],
            "sales_amount": [12, 20, 15],
            "unit_cost": [4, 7, 5],
        }
    )
    frame.to_csv(source, index=False)

    from studio_api.planning import build_plan
    from studio_api.profiling import profile_file

    plan = build_plan(profile_file(source), "哪个产品销售额最高？")

    assert plan.status == "READY"
    assert plan.fields == {"dimension": "product_name", "metric": "sales_amount"}


def test_sales_volume_question_does_not_use_a_sales_amount_field(tmp_path):
    source = tmp_path / "orders.csv"
    frame = pd.DataFrame(
        {
            "product_name": ["A", "B", "A"],
            "sales_amount": [12, 20, 15],
        }
    )
    frame.to_csv(source, index=False)

    from studio_api.planning import build_plan
    from studio_api.profiling import profile_file

    plan = build_plan(profile_file(source), "哪个产品销量最高？")

    assert plan.status == "INSUFFICIENT_DATA"
    assert plan.fields == {"dimension": "product_name"}


def test_question_terms_do_not_fall_back_to_unrequested_business_fields(tmp_path):
    source = tmp_path / "regional.csv"
    frame = pd.DataFrame(
        {
            "month": ["2025-01-01", "2025-02-01", "2025-03-01"],
            "region": ["north", "south", "north"],
            "revenue": [12, 20, 15],
        }
    )
    frame.to_csv(source, index=False)

    from studio_api.planning import build_plan
    from studio_api.profiling import profile_file

    plan = build_plan(profile_file(source), "哪个产品利润最高？")

    assert plan.status == "INSUFFICIENT_DATA"
    assert plan.fields == {}


def test_generic_trend_uses_the_profile_selected_time_and_metric(tmp_path):
    source = tmp_path / "series.csv"
    frame = pd.DataFrame(
        {
            "dt": ["2025-01-01", "2025-01-02", "2025-01-03"],
            "value_x": [3, 5, 4],
        }
    )
    frame.to_csv(source, index=False)

    from studio_api.execution import execute_plan
    from studio_api.planning import build_plan
    from studio_api.profiling import profile_file

    profile = profile_file(source)
    plan = build_plan(profile, "value_x 随 dt 的趋势如何？")
    result = execute_plan({profile.tables[0].name: frame}, plan)

    assert plan.fields == {"time": "dt", "metric": "value_x"}
    assert result.status == "SUCCESS"
    assert result.findings[0].value == [
        {"time": "2025-01-01", "value": 3.0},
        {"time": "2025-01-02", "value": 5.0},
        {"time": "2025-01-03", "value": 4.0},
    ]
    assert result.findings[0].evidence.calculation == "groupby(dt).sum(value_x).sort_index()"


def test_generic_group_comparison_uses_mean_by_requested_dimension(tmp_path):
    source = tmp_path / "scores.csv"
    frame = pd.DataFrame(
        {
            "cohort": ["blue", "blue", "red", "red"],
            "score_value": [8, 12, 15, 25],
        }
    )
    frame.to_csv(source, index=False)

    from studio_api.execution import execute_plan
    from studio_api.planning import build_plan
    from studio_api.profiling import profile_file

    profile = profile_file(source)
    plan = build_plan(profile, "比较 cohort 的 score_value 差异")
    result = execute_plan({profile.tables[0].name: frame}, plan)

    assert plan.operations == ("group_comparison",)
    assert plan.fields == {"dimension": "cohort", "metric": "score_value"}
    assert result.findings[0].value == [
        {"group": "red", "value": 20.0},
        {"group": "blue", "value": 10.0},
    ]
    assert result.findings[0].evidence.calculation == "groupby(cohort).mean(score_value)"


def test_generic_correlation_uses_two_requested_metrics(tmp_path):
    source = tmp_path / "paired.csv"
    frame = pd.DataFrame(
        {
            "value_x": [1, 2, 3, 4],
            "value_y": [2, 4, 6, 8],
        }
    )
    frame.to_csv(source, index=False)

    from studio_api.execution import execute_plan
    from studio_api.planning import build_plan
    from studio_api.profiling import profile_file

    profile = profile_file(source)
    plan = build_plan(profile, "value_x 与 value_y 是否相关？")
    result = execute_plan({profile.tables[0].name: frame}, plan)

    assert plan.operations == ("correlation",)
    assert plan.fields == {"metric": "value_x", "secondary_metric": "value_y"}
    assert result.findings[0].value == 1.0
    assert result.findings[0].evidence.calculation == "pearson_correlation(value_x, value_y)"


def make_sales_workbook() -> bytes:
    workbook = Workbook()
    sheet = workbook.active
    sheet.title = "sales"
    sheet.append(["month", "region", "revenue"])
    for month, revenue in enumerate([100, 120, 150, 180, 220, 260, 310, 360, 420, 490, 570, 660], 1):
        sheet.append([f"2025-{month:02d}-01", "north" if month % 2 else "south", revenue])
    buffer = BytesIO()
    workbook.save(buffer)
    return buffer.getvalue()


def test_dataset_upload_serializes_excel_datetime_profile_values():
    workbook = Workbook()
    sheet = workbook.active
    sheet.title = "dated_metrics"
    sheet.append(["observed_at", "measure_x"])
    sheet.append([datetime(2025, 1, 1), 10])
    sheet.append([datetime(2025, 2, 1), 20])
    buffer = BytesIO()
    workbook.save(buffer)

    response = client.post(
        "/api/datasets",
        files={"file": ("dated-metrics.xlsx", buffer.getvalue(), "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet")},
    )

    assert response.status_code == 201
    profile = response.json()["intake"]["profile"]
    assert profile["tables"][0]["columns"][0]["samples"] == ["2025-01-01T00:00:00", "2025-02-01T00:00:00"]


def make_regional_revenue_workbook(missing_south_march: bool = False) -> bytes:
    workbook = Workbook()
    sheet = workbook.active
    sheet.title = "regional revenue"
    sheet.append(["month", "region", "revenue"])
    north = [100, 103, 107, 110, 113, 116, 120, 123, 126, 130, 134, 138]
    south = [155, 150, 145, 137, 128, 118, 108, 98, 90, 84, 79, 74]
    east = [112, 115, 118, 92, 116, 119, 122, 121, 124, 127, 130, 133]
    for month in range(1, 13):
        for region, revenue in (("north", north[month - 1]), ("south", None if missing_south_march and month == 3 else south[month - 1]), ("east", east[month - 1])):
            sheet.append([f"2025-{month:02d}-01", region, revenue])
    buffer = BytesIO()
    workbook.save(buffer)
    return buffer.getvalue()


def write_financial_workbook(tmp_path):
    workbook = Workbook()
    cover = workbook.active
    cover.title = "封面"
    cover.append(["2025 年度经营分析"])

    sheet = workbook.create_sheet("财务汇总")
    sheet.append(["经营月报"])
    sheet.append(["单位：万元"])
    sheet.append([])
    sheet.append(["期间", "营业收入", "毛利率", "营业利润"])
    sheet.append(["2025-01", 100, 0.32, 15])
    sheet.append(["2025-02", 120, 0.35, 18])

    path = tmp_path / "financial.xlsx"
    workbook.save(path)
    return path


def write_multi_metric_financial_workbook(tmp_path):
    workbook = Workbook()
    sheet = workbook.active
    sheet.title = "经营明细"
    sheet.append(["经营月报"])
    sheet.append(["单位：万元"])
    sheet.append([])
    sheet.append(["期间", "营业收入（万元）", "毛利（万元）", "营业利润（万元）"])
    sheet.append(["2023-12", 9999, 3000, 1200])
    for index, period in enumerate(pd.period_range("2024-01", "2025-12", freq="M"), 1):
        revenue = 100 + index * 10
        primary_margin, secondary_margin = 0.30, 0.20
        operating_profit = revenue * 0.15
        if str(period) == "2025-03":
            revenue = 600
            primary_margin, secondary_margin = 0.20, 0.10
            operating_profit = 2
        primary_revenue = revenue * 0.9
        secondary_revenue = revenue - primary_revenue
        sheet.append([str(period), primary_revenue, primary_revenue * primary_margin, operating_profit * 0.9])
        sheet.append([str(period), secondary_revenue, secondary_revenue * secondary_margin, operating_profit * 0.1])

    path = tmp_path / "multi-metric-financial.xlsx"
    workbook.save(path)
    return path


def write_workbook_with_repeated_detail_title(tmp_path):
    workbook = Workbook()
    detail = workbook.active
    detail.title = "经营明细"
    detail.append(["经营明细", "经营明细", "经营明细", "经营明细"])
    detail.append(["说明", "说明", "说明", "说明"])
    detail.append(["一", "二", "三", "四"])

    sheet = workbook.create_sheet("月度财务")
    sheet.append(["2025 年经营分析"])
    sheet.append(["单位：万元"])
    sheet.append([])
    sheet.append(["期间", "营业收入", "毛利率", "营业利润"])
    sheet.append(["2025-01", 100, 0.32, 15])
    sheet.append(["2025-02", 120, 0.35, 18])

    path = tmp_path / "monthly-financial.xlsx"
    workbook.save(path)
    return path


def write_workbook_with_date_and_text_only_sheet(tmp_path):
    workbook = Workbook()
    schedule = workbook.active
    schedule.title = "日程说明"
    schedule.append(["日期", "事项", "负责人", "备注"])
    schedule.append([datetime(2025, 1, 1), "启动会", "张三", "完成"])
    schedule.append([datetime(2025, 2, 1), "复盘会", "李四", "待定"])

    sheet = workbook.create_sheet("月度收入")
    sheet.append(["月度收入分析"])
    sheet.append([])
    sheet.append(["期间", "营业收入"])
    sheet.append(["2025-01", 100])
    sheet.append(["2025-02", 120])

    path = tmp_path / "date-and-text.xlsx"
    workbook.save(path)
    return path


def analyse_uploaded(objective: str, missing_south_march: bool = False) -> dict:
    create = client.post(
        "/api/jobs",
        data={"objective": objective},
        files={"file": ("regional.xlsx", make_regional_revenue_workbook(missing_south_march), "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet")},
    )
    assert create.status_code == 201
    result = client.post(f"/api/jobs/{create.json()['id']}/analyze")
    assert result.status_code == 200
    return result.json()


def test_excel_job_returns_generic_findings_and_reports():
    create = client.post(
        "/api/jobs",
        data={"objective": "哪个 region 的 revenue 最高？"},
        files={"file": ("sales.xlsx", make_sales_workbook(), "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet")},
    )

    assert create.status_code == 201
    job = create.json()
    assert job["intake"]["kind"] == "spreadsheet"
    assert job["intake"]["columns"] == ["month", "region", "revenue"]

    run = client.post(f"/api/jobs/{job['id']}/analyze")

    assert run.status_code == 200
    result = run.json()
    assert result["status"] == "succeeded"
    assert result["validation_status"] == "SUCCESS"
    assert result["analysis"]["kind"] == "structured_analysis"
    assert result["findings"][0]["evidence"]["calculation"] == "groupby(region).sum(revenue)"
    assert result["findings"][0]["evidence"]["source"]["file_hash"]
    assert {report["format"] for report in result["reports"]} == {"markdown", "html", "docx"}

    markdown_report = next(report for report in result["reports"] if report["format"] == "markdown")
    markdown = client.get(markdown_report["download_url"]).text
    assert markdown.index("## 核心结论") < markdown.index("## 数据质量与分析限制")
    assert "groupby" not in markdown

    word_report = next(report for report in result["reports"] if report["format"] == "docx")
    report = client.get(word_report["download_url"])
    assert report.status_code == 200
    headings = [paragraph.text for paragraph in Document(BytesIO(report.content)).paragraphs]
    assert "核心结论" in headings
    assert "数据质量与分析限制" in headings


def test_read_spreadsheet_selects_a_monthly_table_after_cover_rows(tmp_path):
    frame = read_spreadsheet(write_financial_workbook(tmp_path))

    assert frame.attrs["source_sheet"] == "财务汇总"
    assert frame.attrs["header_row"] == 3
    assert {"期间", "营业收入", "毛利率", "营业利润"} <= set(frame.columns)


def test_read_spreadsheet_skips_repeated_merged_titles_for_valid_monthly_table(tmp_path):
    frame = read_spreadsheet(write_workbook_with_repeated_detail_title(tmp_path))

    assert frame.attrs["source_sheet"] == "月度财务"
    assert frame.attrs["header_row"] == 3
    assert pd.to_datetime(frame["期间"], errors="coerce").notna().sum() == 2
    assert pd.to_numeric(frame["营业收入"], errors="coerce").notna().sum() == 2


def test_read_spreadsheet_rejects_date_and_text_only_sheet_as_a_valid_table(tmp_path):
    frame = read_spreadsheet(write_workbook_with_date_and_text_only_sheet(tmp_path))

    assert frame.attrs["source_sheet"] == "月度收入"
    assert "营业收入" in frame.columns


def test_financial_question_answers_all_requested_metrics_with_monthly_evidence(tmp_path):
    request = "分析 2024-01 到 2025-12 的营业收入、毛利率和营业利润趋势，指出异常月份及可能原因，并引用数据证据。"

    frame = read_spreadsheet(write_multi_metric_financial_workbook(tmp_path))
    plan = plan_question(frame, request)
    result = analyse_spreadsheet(frame, request, tmp_path)

    assert plan.metric_columns == ("营业收入（万元）", "毛利率", "营业利润（万元）")
    assert str(plan.period_start)[:7] == "2024-01"
    assert str(plan.period_end)[:7] == "2025-12"
    assert all(name in result["core_conclusion"] for name in ("营业收入", "毛利率", "营业利润"))
    assert "毛利率 在 2024-01 至 2025-12 呈稳定趋势，从 0.29 到 0.29" in result["core_conclusion"]
    assert "2025-03" in result["core_conclusion"]
    assert "9999" not in result["core_conclusion"]
    assert {item["label"] for item in result["key_metrics"]} >= set(plan.metric_columns)
    assert result["charts"]
    chart = (tmp_path / result["charts"][0]["path"]).read_text(encoding="utf-8")
    assert all(name.encode("unicode_escape").decode("ascii") in chart for name in ("营业收入", "毛利率", "营业利润", "异常"))


def test_financial_anomaly_keeps_the_decline_and_grounds_its_reason_in_workbook_fields(tmp_path):
    frame = pd.DataFrame(
        {
            "期间": ["2025-01", "2025-02", "2025-03", "2025-04"],
            "销量（件）": [970, 1000, 841, 1068],
            "含税单价（元）": [100, 100, 100, 100],
            "营业收入（万元）": [97, 100, 84.1, 106.8],
            "毛利（万元）": [29.1, 30, 25.23, 32.04],
            "营业利润（万元）": [14.55, 15, 12.51, 15.95],
            "备注": ["正常经营", "正常经营", "需求回落：销量下滑", "正常经营"],
        }
    )

    result = analyse_spreadsheet(frame, "分析 2025-01 到 2025-04 的营业收入、毛利率和营业利润趋势，指出异常月份及可能原因", tmp_path)
    anomaly_text = " ".join(item["text"] for section in result["sections"] if section["title"] == "异常对象" for item in section["items"])
    answer_text = f"{result['core_conclusion']} {anomaly_text}"

    assert "2025-03" in answer_text
    assert "营业收入（万元）" in answer_text and "环比下降 15.9%" in answer_text
    assert "营业利润（万元）" in answer_text and "环比下降 16.6%" in answer_text
    assert "备注“需求回落：销量下滑”" in answer_text
    assert "销量（件）从 1,000 到 841，环比下降 15.9%" in answer_text
    assert "含税单价（元）从 100 到 100，环比稳定 0.0%" in answer_text
    assert "数据未提供客户、价格或业务事件字段" not in answer_text


def test_multi_metric_answer_reports_no_month_above_anomaly_threshold(tmp_path):
    frame = pd.DataFrame(
        {
            "期间": pd.period_range("2025-01", "2025-12", freq="M").astype(str),
            "营业收入": [100 + index * 2 for index in range(12)],
            "毛利额": [30 + index * 0.6 for index in range(12)],
            "营业利润": [15 + index * 0.3 for index in range(12)],
        }
    )

    result = analyse_spreadsheet(frame, "分析 2025-01 到 2025-12 的营业收入、毛利率和营业利润趋势，指出异常月份", tmp_path)

    anomaly_items = next(section["items"] for section in result["sections"] if section["title"] == "异常对象")
    assert anomaly_items == [{"text": "未识别到超过阈值的异常月份。"}]


def test_multi_metric_explanation_reflects_same_direction_margin_change(tmp_path):
    frame = pd.DataFrame(
        {
            "期间": ["2025-01", "2025-02", "2025-03", "2025-04"],
            "营业收入": [100, 101, 102, 129.54],
            "毛利额": [30, 30.3, 30.6, 39.02],
            "营业利润": [15, 15.1, 15.2, 19.38],
        }
    )

    result = analyse_spreadsheet(frame, "分析 2025-01 到 2025-04 的营业收入、毛利率和营业利润趋势，指出异常月份及可能原因", tmp_path)

    assert "2025-04 的指标联动显示：营业收入 环比上升 27.0%、毛利率 环比上升 0.4%、营业利润 环比上升 27.5%" in result["core_conclusion"]
    assert "营业收入与营业利润同向，毛利率上升" in result["core_conclusion"]
    assert "毛利率稳定或略升" not in result["core_conclusion"]


def test_multi_metric_explanation_names_a_declining_margin(tmp_path):
    frame = pd.DataFrame(
        {
            "期间": ["2025-01", "2025-02", "2025-03", "2025-04"],
            "营业收入": [100, 101, 102, 129.54],
            "毛利额": [30, 30.3, 30.6, 36.27],
            "营业利润": [15, 15.1, 15.2, 19.38],
        }
    )

    result = analyse_spreadsheet(frame, "分析 2025-01 到 2025-04 的营业收入、毛利率和营业利润趋势，指出异常月份及可能原因", tmp_path)

    assert "营业收入与营业利润同向，毛利率下降" in result["core_conclusion"]
    assert "毛利率上升" not in result["core_conclusion"] and "毛利率稳定" not in result["core_conclusion"]


def test_multi_metric_explanation_preserves_a_small_declining_margin(tmp_path):
    frame = pd.DataFrame(
        {
            "期间": ["2025-01", "2025-02", "2025-03", "2025-04"],
            "营业收入": [100, 101, 102, 129.54],
            "毛利额": [30, 30.3, 30.6, 38.842569],
            "营业利润": [15, 15.1, 15.2, 19.38],
        }
    )

    result = analyse_spreadsheet(frame, "分析 2025-01 到 2025-04 的营业收入、毛利率和营业利润趋势，指出异常月份及可能原因", tmp_path)

    assert "毛利率 环比下降 0.05%" in result["core_conclusion"]
    assert "营业收入与营业利润同向，毛利率下降" in result["core_conclusion"]
    assert "毛利率稳定" not in result["core_conclusion"]


def test_multi_metric_explanation_preserves_a_tiny_declining_margin(tmp_path):
    frame = pd.DataFrame(
        {
            "期间": ["2025-01", "2025-02", "2025-03", "2025-04"],
            "营业收入": [100, 101, 102, 129.54],
            "毛利额": [30, 30.3, 30.6, 38.86044552],
            "营业利润": [15, 15.1, 15.2, 19.38],
        }
    )

    result = analyse_spreadsheet(frame, "分析 2025-01 到 2025-04 的营业收入、毛利率和营业利润趋势，指出异常月份及可能原因", tmp_path)

    assert "毛利率 环比下降 0.004%" in result["core_conclusion"]
    assert "营业收入与营业利润同向，毛利率下降" in result["core_conclusion"]
    assert "毛利率 环比下降 0.00%" not in result["core_conclusion"]
    assert "毛利率稳定" not in result["core_conclusion"]


def test_operating_profit_only_anomaly_has_a_possible_reason(tmp_path):
    frame = pd.DataFrame(
        {
            "期间": ["2025-01", "2025-02", "2025-03", "2025-04"],
            "营业收入": [100, 101, 102, 103],
            "毛利额": [30, 30.3, 30.6, 30.9],
            "营业利润": [15, 15.1, 15.2, 7.6],
        }
    )

    result = analyse_spreadsheet(frame, "分析 2025-01 到 2025-04 的营业收入、毛利率和营业利润趋势，指出异常月份及可能原因", tmp_path)

    anomaly_items = next(section["items"] for section in result["sections"] if section["title"] == "异常对象")
    assert len(anomaly_items) == 1
    assert "营业利润" in anomaly_items[0]["text"]
    assert "2025-04" in anomaly_items[0]["text"]
    assert "%" in anomaly_items[0]["text"] and "可能原因" in anomaly_items[0]["text"]


def test_word_job_is_reported_as_document_evidence_not_measured_data():
    document = Document()
    document.add_heading("项目风险说明", level=1)
    document.add_paragraph("供应延期可能影响交付，需要在本周确认备选方案。")
    buffer = BytesIO()
    document.save(buffer)

    create = client.post(
        "/api/jobs",
        data={"objective": "审阅项目风险并给出低损害方案"},
        files={"file": ("brief.docx", buffer.getvalue(), "application/vnd.openxmlformats-officedocument.wordprocessingml.document")},
    )
    job = create.json()

    run = client.post(f"/api/jobs/{job['id']}/analyze")

    assert run.status_code == 200
    result = run.json()
    assert result["status"] == "succeeded"
    assert result["validation_status"] == "SUCCESS"
    assert result["analysis"]["kind"] == "document_review"
    assert result["evidence"][0]["level"] == "document_statement"
    assert "文档陈述" in result["evidence"][0]["summary"]
    assert result["reports"]


def test_unsupported_file_is_rejected():
    response = client.post(
        "/api/jobs",
        data={"objective": "分析"},
        files={"file": ("notes.txt", b"hello", "text/plain")},
    )

    assert response.status_code == 422
    assert "只支持" in response.json()["detail"]


def test_one_dataset_can_restore_independent_analysis_sessions():
    dataset_response = client.post(
        "/api/datasets",
        files={"file": ("sales.xlsx", make_sales_workbook(), "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet")},
    )

    assert dataset_response.status_code == 201
    dataset = dataset_response.json()

    trend = client.post("/api/sessions", json={"dataset_id": dataset["id"], "objective": "分析营收趋势"})
    risk = client.post("/api/sessions", json={"dataset_id": dataset["id"], "objective": "检测异常订单风险"})

    assert trend.status_code == 201
    assert risk.status_code == 201
    assert trend.json()["dataset_id"] == dataset["id"]
    assert trend.json()["id"] != risk.json()["id"]
    assert trend.json()["messages"][0]["content"] == "分析营收趋势"

    listed = client.get(f"/api/sessions?dataset_id={dataset['id']}")
    assert listed.status_code == 200
    assert {item["id"] for item in listed.json()} >= {trend.json()["id"], risk.json()["id"]}

    renamed = client.patch(f"/api/sessions/{trend.json()['id']}", json={"title": "月度营收趋势"})
    assert renamed.status_code == 200
    assert renamed.json()["title"] == "月度营收趋势"

    copied = client.post(f"/api/sessions/{trend.json()['id']}/copy")
    assert copied.status_code == 201
    assert copied.json()["dataset_id"] == dataset["id"]
    assert copied.json()["messages"] == []
    assert copied.json()["status"] == "ready"

    analyzed = client.post(f"/api/sessions/{trend.json()['id']}/analyze")
    assert analyzed.status_code == 200
    assert analyzed.json()["status"] == "succeeded"
    assert analyzed.json()["validation_status"] == "SUCCESS"
    assert analyzed.json()["notebook_cells"][0]["language"] == "python"
    assert client.get(f"/api/sessions/{risk.json()['id']}").json()["status"] == "ready"

    deleted = client.delete(f"/api/sessions/{copied.json()['id']}")
    assert deleted.status_code == 204
    assert client.get(f"/api/sessions/{copied.json()['id']}").status_code == 404


def test_session_page_returns_requested_slice_without_changing_list_endpoint(tmp_path, monkeypatch):
    from studio_api import store

    monkeypatch.setattr(store, "ROOT", tmp_path)
    uploaded = client.post(
        "/api/datasets",
        files={"file": ("sales.xlsx", make_sales_workbook(), "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet")},
    )
    dataset_id = uploaded.json()["id"]
    created = [
        client.post("/api/sessions", json={"dataset_id": dataset_id, "objective": f"分析营收趋势 {index}"}).json()
        for index in range(3)
    ]

    page = client.get("/api/sessions/page?offset=0&limit=2")

    assert page.status_code == 200
    assert [item["id"] for item in page.json()["items"]] == [created[2]["id"], created[1]["id"]]
    assert page.json()["next_offset"] == 2
    assert page.json()["has_more"] is True
    assert len(client.get("/api/sessions").json()) == 3


def test_session_analysis_persists_datetime_values_from_a_generic_excel_table(tmp_path, monkeypatch):
    from studio_api import store

    monkeypatch.setattr(store, "ROOT", tmp_path)
    workbook = Workbook()
    sheet = workbook.active
    sheet.append(["data_month", "section_id", "total_fee"])
    for month, fee in enumerate([100, 102, 101, 103, 500, 104], start=1):
        sheet.append([datetime(2022, month, 1), "G03213717", fee])
    content = BytesIO()
    workbook.save(content)

    uploaded = client.post(
        "/api/datasets",
        files={"file": ("toll.xlsx", content.getvalue(), "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet")},
    )
    created = client.post("/api/sessions", json={"dataset_id": uploaded.json()["id"], "objective": "检测 total_fee 异常"})

    result = client.post(f"/api/sessions/{created.json()['id']}/analyze")

    assert result.status_code == 200
    payload = result.json()
    anomaly = next(item for item in payload["findings"] if item["kind"] == "anomaly")
    assert isinstance(anomaly["context"]["time"], str)


def test_composite_financial_question_returns_all_requested_metrics_with_monthly_evidence_via_api():
    workbook = Workbook()
    sheet = workbook.active
    sheet.title = "财务月报"
    sheet.append(["期间", "营业收入", "毛利", "营业利润", "备注"])
    for offset, period in enumerate(pd.period_range("2024-01", "2025-12", freq="M")):
        revenue = 100 + offset
        gross_profit = revenue * 0.3
        operating_profit = revenue * 0.12
        note = "正常经营"
        if str(period) == "2025-06":
            revenue = 260
            gross_profit = 78
            operating_profit = 40
            note = "临时大额订单确认"
        sheet.append([str(period), revenue, gross_profit, operating_profit, note])
    content = BytesIO()
    workbook.save(content)
    objective = "分析 2024-01 到 2025-12 的营业收入、毛利率和营业利润趋势，指出异常月份及可能原因，并引用数据证据。"

    uploaded = client.post(
        "/api/datasets",
        files={"file": ("monthly-finance.xlsx", content.getvalue(), "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet")},
    )
    assert uploaded.status_code == 201
    created = client.post("/api/sessions", json={"dataset_id": uploaded.json()["id"], "objective": objective})
    assert created.status_code == 201

    result = client.post(f"/api/sessions/{created.json()['id']}/analyze")

    assert result.status_code == 200
    payload = result.json()
    assert payload["status"] == "succeeded"
    assert payload["validation_status"] == "SUCCESS"
    assert "未找到满足语义置信度 >= 0.70 的必要字段：metric" not in payload["answer"]
    assert {finding["context"]["metric"] for finding in payload["findings"] if finding["kind"] == "trend"} == {"营业收入", "毛利率", "营业利润"}
    assert "2025-06" in payload["answer"]
    assert "可能原因线索" in payload["answer"]
    assert payload["answer"].startswith("结论：")
    assert "趋势：营业收入" in payload["answer"]
    assert "异常月份：" in payload["answer"]
    assert "毛利率：未发现超过 IQR 阈值的月度异常" in payload["answer"]
    assert "原因证据：" in payload["answer"]
    assert all(finding["evidence"]["source"]["table"] == "财务月报" for finding in payload["findings"])
    margin = next(finding for finding in payload["findings"] if finding["kind"] == "trend" and finding["context"]["metric"] == "毛利率")
    assert margin["evidence"]["calculation"] == "monthly_sum(毛利) / monthly_sum(营业收入)"
    assert margin["evidence"]["formula"] == "sum(毛利) / sum(营业收入)"


def test_semantic_session_titles_are_short_and_duplicates_are_numbered():
    dataset_response = client.post(
        "/api/datasets",
        files={"file": ("sales.xlsx", make_sales_workbook(), "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet")},
    )
    dataset_id = dataset_response.json()["id"]
    objective = "分析不同地区营收情况，检测异常地区并预测未来营收风险"

    first = client.post("/api/sessions", json={"dataset_id": dataset_id, "objective": objective})
    second = client.post("/api/sessions", json={"dataset_id": dataset_id, "objective": objective})

    assert first.status_code == 201
    assert first.json()["title"].startswith("地区营收风险")
    assert second.status_code == 201
    assert second.json()["title"].startswith("地区营收风险 ·")
    assert first.json()["title"] != second.json()["title"]


def test_same_question_uses_a_unique_title_across_datasets():
    first_dataset = client.post("/api/datasets", files={"file": ("first.xlsx", make_sales_workbook(), "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet")}).json()
    second_dataset = client.post("/api/datasets", files={"file": ("second.xlsx", make_sales_workbook(), "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet")}).json()
    objective = "检测地区营收异常风险"

    first = client.post("/api/sessions", json={"dataset_id": first_dataset["id"], "objective": objective})
    second = client.post("/api/sessions", json={"dataset_id": second_dataset["id"], "objective": objective})

    assert first.json()["title"] != second.json()["title"]
    assert second.json()["title"].startswith("地区营收风险 ·")


def test_question_plan_recognises_region_revenue_risk_request():
    frame = pd.DataFrame(
        {
            "month": ["2025-01", "2025-02"],
            "region": ["north", "south"],
            "revenue": [100, 80],
        }
    )

    plan = plan_question(frame, "检测地区营收异常风险并预测未来风险")

    assert set(plan.types) == {"anomaly", "risk", "forecast"}
    assert (plan.time_column, plan.metric_column, plan.dimension_column) == ("month", "revenue", "region")


def test_question_plan_prefers_the_requested_gmv_metric_over_identifier_columns():
    frame = pd.DataFrame({"month": ["2025-01", "2025-02"], "region": ["north", "south"], "order_id": [1, 2], "GMV": [100, 80]})

    plan = plan_question(frame, "分析GMV趋势")

    assert plan.metric_column == "GMV"
    assert plan.title == "GMV趋势"


def test_generic_anomaly_response_keeps_executor_evidence():
    result = analyse_uploaded("检测地区营收异常风险")

    assert result["status"] == "succeeded"
    assert result["validation_status"] == "SUCCESS"
    assert result["findings"]
    assert result["findings"][0]["evidence"]["calculation"] == "iqr_outliers(revenue, k=1.5)"
    assert result["findings"][0]["evidence"]["output_value"] == result["findings"][0]["value"]


def test_generic_trend_response_contains_only_traceable_conclusions():
    result = analyse_uploaded("analysis revenue trend")

    assert result["status"] == "succeeded"
    assert result["validation_status"] == "SUCCESS"
    assert result["analysis"]["plan"]["operations"] == ["trend"]
    assert result["findings"][0]["evidence"]["calculation"] == "monthly_sum(revenue)"


def test_unsupported_forecast_is_reported_as_insufficient_data():
    result = analyse_uploaded("预测未来营收")

    assert result["status"] == "succeeded"
    assert result["validation_status"] == "INSUFFICIENT_DATA"
    assert result["findings"] == []
    assert any("forecast" in limitation for limitation in result["limitations"])


def test_generic_answer_is_derived_from_validated_findings():
    result = analyse_uploaded("检测地区营收异常风险")

    assert result["core_conclusion"] != result["data_quality"]["summary"]
    assert result["core_conclusion"] == result["answer"]
    assert result["answer"] == "；".join(item["conclusion"] for item in result["findings"])


def test_generic_trend_uses_the_same_source_for_all_evidence():
    result = analyse_uploaded("分析地区营收趋势", missing_south_march=True)

    assert result["status"] == "succeeded"
    assert result["validation_status"] == "SUCCESS"
    assert all(item["source"]["table"] == "regional revenue" for item in result["evidence"])


def test_forecast_does_not_invent_numeric_output_before_executor_support_exists():
    result = analyse_uploaded("预测未来营收")
    assert result["status"] == "succeeded"
    assert result["validation_status"] == "INSUFFICIENT_DATA"
    assert not result["evidence"]


def test_generic_ranking_answers_the_executor_leader():
    result = analyse_uploaded("哪个地区表现最好？")

    assert "east" in result["core_conclusion"].lower()
    assert result["findings"][0]["metric_value"] == 1429.0
    assert result["findings"][0]["evidence"]["calculation"] == "groupby(region).sum(revenue)"


def test_generic_result_does_not_claim_a_chart_when_no_chart_executor_exists():
    anomaly = analyse_uploaded("检测地区营收异常风险")
    forecast = analyse_uploaded("预测未来营收")

    assert anomaly["charts"] == []
    assert forecast["charts"] == []


def test_future_question_preserves_unsupported_operation_as_a_limitation():
    result = analyse_uploaded("哪些地区未来可能继续下滑")

    assert result["analysis"]["plan"]["operations"] == ["forecast"]
    assert result["status"] == "succeeded"
    assert result["validation_status"] == "INSUFFICIENT_DATA"
    assert any("forecast" in limitation for limitation in result["limitations"])


def test_non_numeric_dataset_returns_a_clear_inability_answer_not_an_internal_error():
    workbook = Workbook()
    sheet = workbook.active
    sheet.append(["month", "region", "comment"])
    sheet.append(["2025-01", "south", "missing revenue"])
    buffer = BytesIO()
    workbook.save(buffer)
    create = client.post("/api/jobs", data={"objective": "分析地区营收趋势"}, files={"file": ("text.xlsx", buffer.getvalue(), "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet")})

    response = client.post(f"/api/jobs/{create.json()['id']}/analyze")

    assert response.status_code == 200
    assert response.json()["status"] == "succeeded"
    assert response.json()["validation_status"] == "INSUFFICIENT_DATA"
    assert "metric" in response.json()["core_conclusion"]


def test_dataset_upload_rejects_a_header_only_csv_with_a_recovery_message():
    response = client.post(
        "/api/datasets",
        files={"file": ("empty.csv", b"month,region,revenue\n", "text/csv")},
    )

    assert response.status_code == 422
    assert response.json()["detail"] == "文件没有数据行，请补充数据后重新上传。"


def test_dataset_upload_does_not_expose_parser_internals_to_users():
    response = client.post(
        "/api/datasets",
        files={"file": ("broken.xlsx", b"not an Excel workbook", "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet")},
    )

    assert response.status_code == 422
    assert response.json()["detail"] == "无法读取该文件，请确认它是可正常打开的 Excel 或 CSV 文件后重试。"

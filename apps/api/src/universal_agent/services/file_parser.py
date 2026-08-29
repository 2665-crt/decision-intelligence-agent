from pathlib import Path

import openpyxl
from docx import Document


def parse_file(path: Path, suffix: str) -> dict:
    if suffix.lower() in {".xlsx", ".xls"}:
        workbook = openpyxl.load_workbook(path, read_only=True, data_only=False)
        sheets = []
        for sheet in workbook.worksheets:
            row = next(sheet.iter_rows(min_row=1, max_row=1, values_only=True), ())
            sheets.append({"name": sheet.title, "columns": [str(value) for value in row if value is not None], "row_count": sheet.max_row})
        return {"sheets": sheets}
    if suffix.lower() == ".docx":
        document = Document(path)
        return {"paragraph_count": len(document.paragraphs), "table_count": len(document.tables), "text_evidence": [{"level": "document_statement", "position": index, "text": paragraph.text} for index, paragraph in enumerate(document.paragraphs) if paragraph.text.strip()]}
    raise ValueError("unsupported file type")

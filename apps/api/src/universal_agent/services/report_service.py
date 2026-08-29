from html import escape
from pathlib import Path

from docx import Document

from universal_agent.services.decision_service import DecisionReport, EvidenceItem, Option, RiskItem


def _evidence_lines(items: list[EvidenceItem]) -> str:
    return "\n".join(f"- {item.summary}" for item in items) or "- 暂无可验证内容"


def _risk_lines(items: list[RiskItem]) -> str:
    blocks = []
    for index, item in enumerate(items, start=1):
        evidence = "；".join(evidence.summary for evidence in item.evidence)
        mitigation = "；".join(item.mitigation)
        blocks.append(
            f"### 风险 {index}：{item.description}\n"
            f"- 概率：{item.probability}\n"
            f"- 影响：{item.impact}\n"
            f"- 严重度：{item.severity}\n"
            f"- 可控性：{item.controllability}\n"
            f"- 证据：{evidence}\n"
            f"- 不确定性：{item.uncertainty}\n"
            f"- 缓解措施：{mitigation}\n"
            f"- 人工复核：{'是' if item.human_review_required else '否'}"
        )
    return "\n\n".join(blocks)


def _option_table(items: list[Option]) -> str:
    header = "| 排名 | 方案 | 预期收益 | 实施成本 | 潜在损害 | 不确定性 | 前提 | 验证指标 | 硬约束满足 |\n|---:|---|---:|---:|---:|---:|---|---|---|"
    rows = [
        f"| {index} | {item.name} | {item.expected_benefit:.2f} | {item.implementation_cost:.2f} | "
        f"{item.potential_harm:.2f} | {item.uncertainty:.2f} | {'；'.join(item.assumptions) or '无'} | "
        f"{item.validation_metric} | {'是' if item.hard_constraints_met else '否'} |"
        for index, item in enumerate(items, start=1)
    ]
    return "\n".join([header, *rows]) if rows else f"{header}\n| - | 暂无可比较方案 | - | - | - | - | - | - | - |"


def _markdown(report: DecisionReport) -> str:
    groups = report.evidence_by_level
    suggestions = "\n".join(f"- {item}" for item in report.suggestions) or "- 暂无"
    return (
        "# 决策支持报告\n\n"
        f"## A级：文件事实\n{_evidence_lines(groups['A'])}\n\n"
        f"## B级：数据推断\n{_evidence_lines(groups['B'])}\n\n"
        f"## C级：预测结果\n{_evidence_lines(groups['C'])}\n\n"
        f"## D级：用户前提与待验证建议\n{_evidence_lines(groups['D'])}\n{suggestions}\n\n"
        f"## 风险登记册\n{_risk_lines(report.risks)}\n\n"
        f"## 方案权衡\n{_option_table(report.options)}\n"
    )


def _add_evidence_section(document: Document, title: str, items: list[EvidenceItem]) -> None:
    document.add_heading(title, level=1)
    if not items:
        document.add_paragraph("暂无可验证内容")
    for item in items:
        document.add_paragraph(item.summary, style="List Bullet")


def _add_risks(document: Document, risks: list[RiskItem]) -> None:
    document.add_heading("风险登记册", level=1)
    for index, risk in enumerate(risks, start=1):
        document.add_heading(f"风险 {index}：{risk.description}", level=2)
        for label, value in (
            ("概率", risk.probability),
            ("影响", risk.impact),
            ("严重度", risk.severity),
            ("可控性", risk.controllability),
            ("证据", "；".join(item.summary for item in risk.evidence)),
            ("不确定性", risk.uncertainty),
            ("缓解措施", "；".join(risk.mitigation)),
            ("人工复核", "是" if risk.human_review_required else "否"),
        ):
            document.add_paragraph(f"{label}：{value}")


def _add_options(document: Document, options: list[Option]) -> None:
    document.add_heading("方案权衡", level=1)
    for index, option in enumerate(options, start=1):
        document.add_heading(f"{index}. {option.name}", level=2)
        document.add_paragraph(
            f"预期收益：{option.expected_benefit:.2f}；实施成本：{option.implementation_cost:.2f}；"
            f"潜在损害：{option.potential_harm:.2f}；不确定性：{option.uncertainty:.2f}"
        )
        document.add_paragraph(f"前提：{'；'.join(option.assumptions) or '无'}")
        document.add_paragraph(f"验证指标：{option.validation_metric}")
        document.add_paragraph(f"硬约束满足：{'是' if option.hard_constraints_met else '否'}")


def render_reports(report: DecisionReport, output_dir: Path) -> list[Path]:
    output_dir.mkdir(parents=True, exist_ok=True)
    markdown = _markdown(report)
    md_path = output_dir / "report.md"
    html_path = output_dir / "report.html"
    docx_path = output_dir / "report.docx"
    md_path.write_text(markdown, encoding="utf-8")
    html_path.write_text(
        f"<!doctype html><html lang='zh-CN'><meta charset='utf-8'><title>决策支持报告</title><body><pre>{escape(markdown)}</pre></body></html>",
        encoding="utf-8",
    )

    document = Document()
    document.add_heading("决策支持报告", level=0)
    groups = report.evidence_by_level
    _add_evidence_section(document, "文件事实", groups["A"])
    _add_evidence_section(document, "数据推断", groups["B"])
    _add_evidence_section(document, "预测结果", groups["C"])
    _add_evidence_section(document, "用户前提", groups["D"])
    document.add_heading("待验证建议", level=1)
    for suggestion in report.suggestions:
        document.add_paragraph(suggestion, style="List Bullet")
    _add_risks(document, report.risks)
    _add_options(document, report.options)
    document.save(docx_path)
    return [md_path, html_path, docx_path]
